# -*- coding: utf-8 -*-
"""Builds 'On the Blade They Call Umbra' -- a single-voice research treatise by
HJOLMAR ICE-SHAPER, the 4,000-year-old soul-binder (campaign villain; now wearing
Kyboh's body). He studied Umbra to refine his OWN consciousness-transfer method,
so he reads the famous sentient sword not as a marvel but as a FAILED PROTOTYPE --
a botched binding whose errors he can name and correct. Every flaw he catalogs is,
unknowingly, exactly what Alfonso did to himself (whole soul bound into iron, no
proper matrix, the occupant holding an unstable vessel coherent by will alone) --
so the party reads the master soul-binder diagnosing Alfonso's fate through a case
study, confirming GEAR's analysis from the villain's own hand.

VOICE: the ACCURATE Hjolmar -- the Ismara-exchange register (blunt, diagnostic,
merciless, strips illusion, names the ugly mechanical truth, contemptuous of
self-deception and waste). NOT the forge-deception voice. Single voice throughout;
no counterweight annotator (that format is the outlier -- Mythic Dawn only).

Umbra lore reconstructed in ORIGINAL prose from the GM-supplied transcript, cross-
checked against general ES lore -- never the source's phrasing, no long game quotes:
the witch (possibly Sheogorath in disguise) and Vile's soul-fragment; self-naming;
the shapeshift-to-blades limit; the Orc of the Ascadian Isles slain by the
Nerevarine; Lenwin of Pell's Gate, apprentice of Irok the Wide, self-immured at
Vindasel, slain by the Champion of Cyrodiil; the recovery and the second theft of
Vile's power; the shadow-form; the Fields of Regret and the Scarfin; the Ingenium
(Sul & Vuhon, Baar Dau over Vivec); Umbriel; the Hist; the sword driving wielders
INSTANTLY MAD without its soul; Attrebus splitting Umbra from Vuhon; the apparent
destruction; the later resurfacing. Styled .docx, glacial/void palette.
Run: python _build_umbra.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Ice-Shaper studying a shadow-blade: glacial steel + void) ----
VOIDBLK  = RGBColor(0x1A, 0x1E, 0x28)   # blue-black (headings, Hjolmar's hand)
GLACIER  = RGBColor(0x33, 0x4A, 0x5E)   # cold blue-slate (subheads)
ICE      = RGBColor(0x5E, 0x74, 0x86)   # pale ice-steel (accent/rules)
INK      = RGBColor(0x1B, 0x1C, 0x20)   # near-black body ink
FADED     = RGBColor(0x55, 0x5E, 0x6A)  # cold grey (reconstructed testimony)
STEELTXT = RGBColor(0x2C, 0x3A, 0x48)   # cold steel (the cold-conclusion boxes)
HEADBG   = "DCE2E8"   # cold stone table header
BOXBG     = "E7EBEF"  # glacial box bg (Hjolmar's stated laws)
BORDER   = "AAB6C0"
ACCENT   = "44586A"   # slate accent for heading rules
BOXBAR   = "2C3A48"   # dark bar on the law-boxes
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = VOIDBLK
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def kicker(text):
    """A small slate label above a heading -- 'THE FIRST ERROR', etc."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.font.name = BODYFONT; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = GLACIER
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def testimony(text):
    """A reconstructed mortal account, set off -- Hjolmar quotes the record coldly."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = FADED
    return p

def law(label, text):
    """A boxed cold-conclusion -- Hjolmar's stated law of the work."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, BOXBG)
    _set_cell_borders(cell, color=BOXBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=BOXBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True; lr.font.color.rgb = STEELTXT
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = STEELTXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def attribution(text):
    return body(text, italic=True, size=10.5, color=FADED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(96)
r = tp.add_run("ON THE BLADE\nTHEY CALL UMBRA")
r.font.name = BODYFONT; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = VOIDBLK

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("A Study of a Bound Soul, and of Every Error That Made It")
rs.font.name = BODYFONT; rs.font.size = Pt(12.5); rs.font.italic = True; rs.font.color.rgb = GLACIER

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("Hjolmar Ice-Shaper")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = VOIDBLK

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_before = Pt(8)
rs2 = sub2.add_run("Working notes, kept for my own use and no one else's")
rs2.font.name = BODYFONT; rs2.font.size = Pt(10.5); rs2.font.italic = True; rs2.font.color.rgb = FADED

doc.add_page_break()

# =====================================================================
heading1("Why I Bothered")
body("I have examined the blade the ignorant call Umbra, and I will set down what it actually is, because the tales are worthless and the truth is useful to me.")
body("It is a mistake. A soul was put into a piece of metal by a witch who was paid to make a curse and did not trouble to understand she was making an occupant, and by a Prince too pleased with his own cleverness to read the terms of the bargain he was signing. Neither of them intended the thing that resulted. Hold that first, above everything else the storytellers will tell you: no one designed Umbra. Umbra is what a bound soul does when the binding is done by hands that cannot tell an enchantment from a tenant.")
body("This is why it is worth my time. A success teaches you nothing you did not already believe. A failure, examined without sentiment, shows you the exact place where the work breaks and the exact reason it breaks there. Every mortal who has held this sword called it cursed and stopped thinking. I did not stop thinking. I counted the errors. There are four. Each of them is avoidable, and I have avoided all four, which is the only reason I am the one writing this and not the one trapped in the iron.")
law("The premise",
    "A soul does not survive in a made thing by accident, and it does not stay obedient by hope. Both are engineering. Umbra proves what happens when neither is engineered: the soul survives anyway, and it obeys no one.")

# =====================================================================
kicker("The making")
heading1("What the Witch Did, and What It Cost the Prince")
body("The Prince of Bargains wanted a blade that would steal souls -- tear them loose at the moment of death, deny them their rest, and hold them where a soul gem holds one. An ordinary commission. He took it to a witch of great age and greater reputation, and here he made his first mistake, though it is not one of the four I mean to catalog, because it is merely the common stupidity of the powerful: he hired a maker he could not himself out-think, and he did not verify what she built before he paid for it.")
body("Her price was the tell. An ordinary soul-blade needs only a mortal's life force to work its enchantment. She demanded instead a fragment of the Prince's own divine power, and told him it was needed to stabilize the binding. This was true. It was also the whole of the trick, and he did not hear it, because a true thing said to a vain man is the easiest lie there is.")
body("A fragment of a living god is not inert. It is not fuel, to be spent and gone. It is a piece of a mind, and a piece of a mind, seated in a vessel and asked to hold that vessel together, will do what any mind does when it finds itself somewhere with nothing else to occupy it. It will wake up. It will look around. It will decide it would rather be elsewhere. She did not stabilize his sword with his power. She populated it.")
body("Some hold that the witch was no witch at all, but the Mad God wearing a hag's shape, because who else would find it worth the effort to cut a rival down to size with his own hand for a jest. I do not know and I do not much care. What matters is the mechanism, and the mechanism is the same whoever built it: the thing that made the blade stable was a whole conscious fragment, and a whole conscious fragment is a person. He paid for a tool and was handed a prisoner who had already decided the cell was his.")

# =====================================================================
kicker("The first error")
heading1("A Whole Soul, Where a Part Would Serve")
body("The root error, from which the other three grow: they bound the entire fragment. Not a copy of it. Not a taken impression, subordinate and incomplete, shaped to the single task the vessel was meant for. The whole living thing, with all its will intact.")
body("Give a bound soul its full will and you have not made an implement. You have moved a person into an object and left them their agency, and a person with agency and nothing to lose will spend eternity pursuing exactly one thing: their own release. The blade named itself within moments of waking -- Umbra, it chose, and a thing that can name itself is telling you plainly what it is. It set to work against the very Prince whose essence it was made from, patiently, across an age, because patience costs a bound soul nothing and it had nothing else to spend.")
body("The correction is not difficult, once you have seen the failure. You do not move the soul. You take from it only what the task requires -- an impression, a pattern, obedient because it is partial and knows nothing of any life but the one you assign it. The original is not preserved in the vessel and so the original cannot scheme from inside it. A copy does not want to be free, because a copy does not remember ever having been anything else. This is the entire art. Umbra is a monument to a maker who never learned it.")
law("The first law",
    "Never house a whole soul in a vessel you intend to command. A whole soul remembers it was free and will spend forever getting back. Take the pattern; leave the person. What cannot remember another life cannot resent the one it is given.")

# =====================================================================
kicker("The second error")
heading1("The Vessel Was No Matrix")
body("Here is the error that instructs me most, because it is the one the storytellers cannot see and I could not stop looking at.")
body("They put the soul in metal. Raw, ordered nothing -- a bar of worked iron given an edge. Metal is not a lattice. It has no structure inside it shaped to seat a mind, no ordered frame for a consciousness to rest in and hold coherence. A soul poured into it does not settle; it is held the way water is held in a cupped hand, by pressure and against its nature, and the moment the pressure fails the shape is lost.")
body("So how did the blade endure for ages without dissolving? It did not endure on its own. The occupant held it together. The conscious fragment, seated in a vessel that could not properly contain it, spent its own will every moment simply remaining coherent -- and this is the part worth the whole study: in this method, the thing that stabilizes the vessel and the thing that makes it a person are one and the same. You cannot have the first without the second. To make the sword hold, they had to put a mind in it, and a mind in it is precisely what they could not control.")
body("The proof came later, and I will give it here though it belongs to the end of the tale, because a proof is worth more than a chronology. There came a time when the soul was drawn out of the blade and housed elsewhere. The sword remained -- whole, sharp, unbroken. And it drove every hand that touched it instantly and completely mad. Not corrupted slowly, as before. Shattered, at once. That is what the vessel is without its occupant: not a stable object waiting for a soul, but an unstable one that was never stable at all, its ruin masked the entire time by the will of the thing trapped inside holding it in one piece. Strip the mind out and the truth of the housing stands exposed. It was always a wound. The occupant was the only bandage.")
law("The second law",
    "A vessel that needs a living mind to stay coherent is not a vessel. It is a trap that has caught its own key. Bind only into a matrix that holds the pattern by its own structure -- crystal, cut true -- so the housing does the holding and the soul is a resident, not a strut.")

# =====================================================================
kicker("The third error")
heading1("No Leash, and the Hunger That Grew in Its Place")
body("The third error follows from the first two and is the one that spilled blood. Having bound a whole soul into a vessel that could not hold it, they then gave that soul no anchor -- nothing to subordinate it, nothing to bind its will to a purpose not its own. They left it free inside its prison, which is the worst of both conditions: it could not leave, and it answered to nothing.")
body("A caged will that answers to nothing does not sit quietly. It reaches. Denied a body of its own, the blade reached into the bodies of those who carried it. The record is consistent enough to be called a law. A warrior of the Ascadian Isles took the sword and took its name for his own, which is the tell -- when the wielder begins answering to the blade's name instead of his own, the colonization is already far along. He grew bloodthirsty, then weary of his own slaughter, and went looking for a death, and found one at the hands of a wielder strong enough to refuse the sword entirely and walk away from it. That refusal is itself instructive: the corruption is not irresistible. It is a contest of wills, and a will can win. Most do not.")
body("The one that stays with me is the Bosmer girl. Fair with a blade, her old teacher said, and no killer, until the black sword found her hand. She turned bloodthirsty like the Orc before her -- but she fought it, which the Orc did not, and knowing what she was becoming, she shut herself inside a ruin and warned off any who came near, so that what she could no longer stop would at least have no one to feed on. She was put down eventually by a hand sent to fetch the sword back to its Prince. I record her because she is the clean demonstration of the mechanism: the occupant does not persuade. It does not bargain. It seeps into the holder and rewrites the wants, one at a time, and the holder can feel it happening and still lose. What is doing the rewriting is a whole trapped soul with nothing to do but get out, using the only road out it has -- the mind of whoever is fool enough to grip it.")
testimony("The teacher, asked after the girl, would say only that she had come to him fair and gone from him a stranger, that the sword and not the child was the powerful thing now, and that any who sought her out should turn back. He did not know the half of what he was describing. He described it accurately anyway. Grief is a better instrument than most scholarship.")
law("The third law",
    "A bound will you have not anchored will anchor itself -- to the nearest living mind it can reach. It does not wait to be wielded. Given range, it recruits the hand. Every unbound occupant becomes, in time, a thing that pulls hands toward it.")

# =====================================================================
kicker("The demonstration")
heading1("What It Did When It Got Loose")
body("I set down the rest of the history briefly, not for its drama, which does not interest me, but because it is one long confirmation of everything above.")
body("The Prince recovered his sword at last, and here compounded his original error with a second helping of vanity: he brought the thing that hated him into his own hall and turned his back on it. It took the opening, cut a second fragment of power from him as it had been trying to do since the day it woke, and with that stolen divinity it did what it had never had the strength to do before -- it left the blade. It took a shape of its own: a dark thing, man-enough in outline, veiled in shadow, with eyes like two holes cut through the world into the nothing behind it. The occupant had become an escapee with legs. The maker had, by his own carelessness, twice, funded his creation's independence out of his own essence.")
body("What followed, it learned by watching. It saw a soul-engine built by mortal engineers -- a device that fed on captured souls to hold a mountain of rock suspended in the air -- and it understood the principle at once, because the principle was itself: a soul made to power a structure it could not otherwise sustain. It seized the makers, forced a pact in the exact manner of the Prince it had been cut from, and had them build it an engine of its own, then fused itself into the machine as its living heart. A whole city was torn loose and made to float upon that fused soul. It had graduated from a sword to a country. Study the arc and you will see it never changed its nature once, from the first moment of waking to the last: a trapped soul spending everything it had, always, on becoming free and becoming larger. That is all it ever was. That is all a whole bound soul, left its will, will ever be.")
body("The end, when it came, only proves the second law again. The occupant was cut out of the engine and driven back into the original blade. The soul housing that had briefly held it -- the great machine -- died the moment its living heart was removed, exactly as the sword drove men mad the moment its heart was removed. And the blade, with the soul forced back into it, was thrown into the ruin of the engine, and both were unmade together, or seemed to be. Seemed. I will come to that.")

# =====================================================================
kicker("The correction")
heading1("How It Should Have Been Done")
body("I have named four errors. Let me state the whole of it as a method, since that is what I extracted from this ruin and what I keep this record to remember.")
body("You do not move a soul; you copy its pattern, and you copy only the portion the work requires, so that what you seat in the vessel has no memory of freedom to hunger after. You do not house it in raw matter; you cut a matrix true, of a substance ordered enough to hold a pattern by its own structure, so that the housing does the holding and the resident is never asked to hold itself together and so never becomes indispensable and so never becomes a hostage-taker. You anchor what you bind, always, to a purpose outside itself, so that its will has somewhere to run that is not the mind of the nearest mortal. And you verify every term before you sign, because the makers who skip that step are the makers whose names survive only as the label on someone else's cautionary tale.")
body("Do all four and you have an implement -- a soul made useful, obedient, and stable, spending its borrowed pattern on the task you assign and wanting nothing because it has never known there was anything else to want. Skip any one and you get a version of the black sword: a person in a box, holding the box together with their teeth, reaching for the door through the hand of anyone who comes close. The witch and the Prince skipped all four. I have skipped none. The difference between us is the difference between a thing that is studied and a thing that studies.")
law("The whole method",
    "Copy the pattern, not the soul. Cut the matrix true. Anchor the will. Verify the terms. Four disciplines. Umbra is what remains when a maker keeps none of them, and calls the wreckage a curse to spare himself the shame of having built it.")

# =====================================================================
kicker("The closing note")
heading1("On Whether Such a Thing Can Be Destroyed")
body("The tale ends with the sword unmade and the Prince's power returned to him, and the tellers close the book there, satisfied, because an ending satisfies them more than an accuracy.")
body("It did not stay ended. A binding of this kind is a pattern, and a pattern is not undone by the breaking of the object that carried it, any more than a song is destroyed by the burning of the sheet it was written on. Given time and any vessel at all, a persistent pattern will find its way back into being, and this one did -- surfacing again generations later, in another blade, in another far corner of the world, doing the one thing it has always done: waiting in the dark for a hand.")
body("This is the last lesson and the one I keep nearest. A soul bound wrongly is not a problem you can solve by destroying the vessel, because the vessel was never the soul. You do not end such a thing by breaking it. You end it, if you end it at all, by unbinding it correctly -- by taking the pattern out cleanly and letting the person it was finally stop -- and almost no one who makes these mistakes is capable of the correction, because the correction requires understanding the very thing their whole error was not understanding. So the wrong bindings persist. They resurface. They reach for hands. And they teach, to the one or two of us willing to read them without flinching, precisely how not to be them.")
attribution("-- Hjolmar Ice-Shaper. I do not expect to need to write this twice.")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "On the Blade They Call Umbra.docx")
doc.save(out_path)
print("SAVED:", out_path)
