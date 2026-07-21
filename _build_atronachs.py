# -*- coding: utf-8 -*-
"""Builds 'Wren Calder's Notebook' -- a bestiary/catalog of Atronachs and
golems, framed as a genuine student research notebook, centuries old, later
adopted as the standard College teaching text on the subject because no
formal treatise since has organized it as clearly. Voice: young, curious,
self-correcting, dated entries -- NOT a settled professor's treatise (see
'On the Negotiable Law' / 'On the Far Side of the Wheel' for that register).
Content drawn from general, verifiable Elder Scrolls lore (elemental Daedra
vs. golems; Infernus/Coldharbour cold-flame variant; Corvus & Calani Direnni
as first binders; the Ash Guardians of Red Mountain; Umbriel-era Xivilai
Scoria) in original prose, not the source video's phrasing.
Styled .docx, warm ink/parchment palette (a notebook, not a lecture hall).
Run: python _build_atronachs.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (a notebook: warm ink-brown, not a lecture-hall palette) ----
INKBROWN = RGBColor(0x4A, 0x30, 0x1E)   # warm brown ink (headings, Wren's hand)
RUST     = RGBColor(0x7A, 0x42, 0x24)   # lighter rust-brown (subheads)
FADED    = RGBColor(0x8A, 0x74, 0x60)   # faded-ink accent (quote/citation ink)
INK      = RGBColor(0x1C, 0x1C, 0x1E)
AMBER    = RGBColor(0x6B, 0x4A, 0x1E)
HEADBG   = "EAE0D0"   # aged-paper table header
MARGBG   = "F0EAE0"   # warm parchment (archivist note bg)
BORDER   = "C7B8A0"
ACCENT   = "8A5A2E"
AMBERBAR = "6B4A1E"
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = INKBROWN
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def entry_date(text):
    """A small dated-entry marker, like a notebook date line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.bold = True
    r.font.color.rgb = RUST
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def struck(text):
    """A self-corrected line -- shown as a crossed-out aside in faded ink."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = FADED
    r.font.strike = True
    return p

def margin_thought(text):
    """A short, informal aside -- Wren talking to herself in the margin."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = RUST
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = FADED
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=FADED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def archivist_note(label, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=AMBERBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=AMBERBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = AMBER
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = AMBER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=INKBROWN)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(90)
r = tp.add_run("A NOTEBOOK ON ATRONACHS AND\nTHEIR FALSE KIN")
r.font.name = BODYFONT; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = INKBROWN

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("Being the Working Notes of an Apprentice, Kept in the\nHope of Someday Understanding What She Was Looking At")
rs.font.name = BODYFONT; rs.font.size = Pt(12); rs.font.italic = True; rs.font.color.rgb = RUST

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(36)
ra = attrib.add_run("Wren Calder\nApprentice, College of Winterhold")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.italic = True; ra.font.color.rgb = INKBROWN

doc.add_page_break()

archivist_note("Archivist's Note -- College of Winterhold, present day",
    "This notebook is centuries old and was never intended for publication; Wren Calder herself appears in no other College record we have located, and we do not know what became of her after these pages stop. It is assigned reading in Magister Vael's own Conjuration curriculum today for a simple reason: no formal treatise written since has organized this subject half as clearly, and her mistakes -- crossed out but left legible -- teach the taxonomy better than a clean answer would. We have reproduced the notebook as found, corrections and all. -- Recordkeeper's Office"
)

# =====================================================================
heading1("The First Page")
entry_date("Sundas, first week of term")
body("Magister says I am to keep a proper notebook this year instead of loose pages, on the theory that loose pages are how I lost six weeks of notes on ward theory last spring. He is right and I resent him for it. So: a notebook. I am going to use it to sort out something that has bothered me since my very first lecture on Conjuration, which is that nobody can actually tell me what an Atronach is.")
body("Not in the way they can tell me what a wolf is, or a Falmer, or even a ghost. Ask three Journeymen what an Atronach is and you get three different answers, and at least one of them will be wrong in a way the other two don't notice. So I am going to work it out for myself, properly, one creature at a time, and write down what I actually find instead of what I was told to expect. If I am still confused by the end of this notebook I will at least be confused in an organized fashion.")
margin_thought("Starting with the flame ones. Everyone starts with the flame ones. There is probably a lesson in that too but I don't have time for it this week.")

# =====================================================================
heading1("The Flame Atronach")
entry_date("Second week")
body("Easiest of the lot to find a live one to study, since half the Destruction faculty keep one on hand for demonstrations. First thing worth writing down: the sketches in the library do not agree with each other, and I don't think either generation of artist was wrong. Older records -- a full century back, some further -- show flame Atronachs with something like a body under the fire: man-shaped, or nearly, wrapped in flame the way a torch is wrapped in its own light. The more recent sketches, mine included, show something with no body to speak of at all. Just fire, holding a shape, with nothing underneath it that isn't also fire.")
body("I don't know why the difference. My first guess was that the older artists were just being polite about how frightening a thing made purely of fire actually is, and gave it a more familiar shape to make it easier to draw. My second guess, which I like better, is that we are looking at the same kind of creature at different points in its own long life, the way a sapling and an old oak are the same kind of thing without looking much alike. I have no way to test this yet. Writing it down so I remember to ask someone who would know.")
body("The part that surprised me: they do not come from any Daedric Prince's realm. I assumed -- everyone assumes -- that anything you summon from Oblivion answers to one Prince or another. Flame Atronachs don't. They have their own small pocket of Oblivion, built collectively by their own kind rather than granted by anyone above them, called Infernus. I asked Magister Vael to describe it and he said, with what I am fairly sure was real feeling, that it is exactly what it sounds like: heat, and rock gone liquid, as far as anything can see. Not designed. Just true to what lives there.")
quote("A demiprince of Boethiah is recorded describing Infernus as \"a dull place\" -- his own words, not mine -- and attributing its plainness to the fact that a realm shaped collectively by many minds of the same kind will always be simpler than one shaped by a single powerful will. I find this genuinely interesting, whatever tone he meant it in. A realm built by consensus tells you something true about what every one of its builders wants, with nothing left over to be personal or strange. Infernus is plain because every fire Daedra who helped build it wanted exactly the same thing, and got it.")

# =====================================================================
heading1("The Cold Flame Atronach")
entry_date("Same week, later")
body("Learned about this one entirely by accident, cross-referencing something unrelated in the Restricted Section and finding a dossier I was not, strictly, supposed to be reading. I am writing it down anyway because it answers a question I would otherwise have spent months on.")
body("There is an inverted flame Atronach -- cold instead of hot, blue flame instead of orange, and by every behavioral account otherwise identical: same temper, same tactics, same habit of hurling fire (cold fire, in this case) at anything that startles it. The dossier belonged to a Coldharbour archivist cataloguing potential menials for Molag Bal's realm, who apparently went looking specifically for a fire-natured servant that could tolerate an ice-cold domain without dying of the mismatch, and searched an enormous number of candidate realms before finding one built on the exact inversion of Infernus. He named it something I have not been able to verify independently, so I am not going to guess at the spelling here and get it permanently wrong in my own notebook.")
margin_thought("The archivist's methodology was thorough. His reasons for it were foul. Worth remembering I can admire the one without excusing the other.")
body("The useful lesson for my actual taxonomy: this proves flame and its opposite are two expressions of the same underlying kind of creature. Keep that in mind for frost.")

# =====================================================================
heading1("The Frost Atronach")
entry_date("Third week")
struck("Assumed, going in, that frost would simply be the cold-flame Atronach's cousin, or possibly the same thing under a different name. Wrong on both counts, and I should have checked before writing that assumption down as though it were settled.")
body("Frost Atronachs turned out to be their own kind entirely -- built of ice rather than of inverted flame, heavier, slower, closer in temperament to a wall that has decided to be angry than to a living flame given a grudge. No one at this College has located their home realm. I asked three separate instructors and got three separate shrugs, which for this subject is apparently the correct scholarly response rather than a failure of the instructors.")
body("What I can say with confidence: frost and flame Atronachs genuinely do not get along, and the friction runs deeper than professional rivalry. Magister Vael described watching the two types deployed side by side during an old account of the Battlespire's fall, forced into uneasy cooperation only because Mehrunes Dagon himself was directing the assault and neither dared refuse him. Left to their own inclinations, apparently, they would rather not share a battlefield at all. I find this more informative than any lecture on their elemental properties -- creatures that cannot stand each other's company are usually creatures whose entire nature is built around a very narrow set of conditions, and both of these plainly are.")

# =====================================================================
heading1("Storm Atronachs, and a Digression on Names")
entry_date("Fourth week")
body("Storm Atronachs I have only seen once, at a demonstration I was not supposed to be close enough to watch. Vaguely man-shaped like the others, but built from something closer to a contained thunderhead than a solid body -- stone and rock-shards suspended in a churn of dark cloud, blue lightning threading through the whole mess without ever quite discharging. Their home realm is named Levinis, so far as the one source I trust on this is willing to commit to paper, though beyond the name I have found nothing else written about the place at all. I would like to know more. I am not in a position to go find out more.")
body("Separately -- and I want to flag this because it nearly derailed my whole taxonomy -- I found a fragment of the Song of Pelinal describing something called a \"fundinark\" fighting alongside Umaril's forces. I spent an embarrassing number of hours convinced this was an old name for the storm Atronach before admitting to myself that I have no actual evidence for that, only a similar-sounding word and my own eagerness to have found something clever. Leaving the guess in the notebook because the process of being wrong about it taught me more than being right would have. The word Atronach itself, incidentally, seems to get borrowed loosely across a lot of old sources for anything animated from raw material rather than born -- which is, I am starting to suspect, exactly the confusion at the root of everything I am trying to sort out in this notebook.")

# =====================================================================
heading1("What I Am Actually Confused About")
entry_date("Fifth week -- the turn")
body("I need to stop cataloguing individual types for a moment and write down the actual problem, because I have been circling it for a month without naming it plainly.")
body("Golems and Atronachs get called the same thing constantly, and they are not the same thing. An Atronach is Daedric -- summoned from Oblivion, a genuine citizen of one of those elemental pocket-realms, brought here rather than made here. A golem is built. A mortal takes some material -- ash, iron, flesh, stone -- and a mortal's own craft animates it. No realm of Oblivion involved at all. The confusion happens because several golem types are built from the same raw elements the Atronachs are made of, so a stone golem and a stone Atronach look, to someone who has not thought about it carefully, like they ought to be cousins. They are not cousins. One is a person from somewhere else. The other is a very sophisticated tool.")
body("Once I wrote that distinction down in one place, most of my earlier confusion resolved on its own. I am annoyed it took me five weeks.")

# =====================================================================
heading1("The Ash Guardians")
entry_date("Sixth week")
body("First golem I have found worth a full entry, and a good test of the distinction I just worked out. Ash Guardians are Telvanni work, built from heartstones -- fragments, so the account I trust says, flung from Red Mountain during its eruption, still carrying a residue of heat from that event centuries later. The theory I have seen argued most convincingly is that the heartstones sat near enough to Lorkhan's own remains for long enough that something of his lingering power worked its way into the rock, and that is what actually animates the ash into a Guardian rather than leaving it as very warm gravel.")
body("If that account is correct -- and I have no strong reason to doubt it, though I have also not verified it myself -- then Ash Guardians are about as far from a Daedric Atronach as a construct can get. They are made from a dead god's own residue, not summoned from any Prince's territory. Calling them Daedric, which I have seen done in at least two texts now, is precisely the kind of careless labeling I keep tripping over.")

# =====================================================================
heading1("The Flesh Atronach")
entry_date("Seventh week")
body("Despite the name, also a golem and not a true Atronach, by my own definition above -- and the most unsettling entry so far, which I am recording plainly rather than dressing up.")
body("Credited to a Dunmer researcher whose work got her expelled from the Mages Guild before it found a home, unsurprisingly, in the Shivering Isles. Her method: stitching together salvaged flesh and bone into a working, animate whole, marked with warding sigils and collared at the throat. Unlike stone or ice, flesh can apparently be shaped with real precision once a mage has the stomach for the work -- built tall, built armed, built for a specific task, in a way rock and ice simply can't be finessed. I have seen it argued that what animates the finished construct is a captured Daedric vestige rather than any mortal soul, worked into the flesh the way a soul is worked into a soul gem. I don't fully trust my own understanding of that last point yet and want to read further before I commit it to this notebook as settled.")
margin_thought("Come back to this entry once I have actually read the source material on vestiges properly. Do not write down secondhand metaphysics as though I checked it myself. I did that with the storm Atronach name and it cost me a week.")

# =====================================================================
heading1("A Warning, Recorded Because I Was Told To")
entry_date("Eighth week")
body("Magister Vael asked me to include this account specifically, and I am including it as told rather than in my own words, because it is not really about taxonomy and I don't want to soften it by making it sound like one of my own entries.")
body("During his campaign against the Empire, Mehrunes Dagon attempted to strengthen one of his Xivilai commanders -- a warrior named Scoria -- by forcing a fusion between the Xivilai's own vestige and the nature of a flame Atronach. The result, by every account, was a genuinely new kind of monster: vastly more powerful, and permanently, agonizingly bound to the conditions a flame Atronach requires to survive. Scoria cannot leave the presence of magma without beginning to fail, and is described, in the one account I have read of it, as suffering something like a continuous burning it can never actually escape, because it has become the thing that burns.")
quote("The source I have for this is secondhand -- an account attributed to a servant of Molag Bal, which is not a witness I would ordinarily trust without a second source to weigh against it. I am recording it anyway, exactly as given to me, because Magister Vael considered it important enough to interrupt my own research to make sure it was in this notebook, and I have decided that is reason enough.")
body("Scoria doesn't fit the taxonomy I've built -- some third thing the fusion made, sitting outside both categories I've spent this whole notebook sorting. Possibly the honest answer is that forcing two kinds of creature together at the level of whatever animates them produces a wound wearing a shape, and calls it a weapon. I am leaving this entry without a tidy conclusion, because I don't have one, and I would rather admit that than invent one.")

# =====================================================================
heading1("Closing This Volume")
entry_date("End of term")
body("I set out to answer one question -- what is an Atronach, actually -- and I think I have an answer now, even if it took a whole term and several wrong guesses to get there. An Atronach is a citizen of Oblivion, elemental and Daedric, belonging to a realm its own kind built collectively rather than one granted by a Prince. A golem is a mortal's construction, animated by craft and, in at least one case I have now studied closely, by something considerably stranger than craft. Everything that confused me in the first week traces back to people using one word for both, and I no longer think that confusion is stupidity on anyone's part. I think it is just what happens when two different kinds of thing are built from the same raw materials and nobody stops to ask what actually gives either one its life.")
body("I am aware this notebook is full of things I got wrong before I got them right, left crossed out rather than erased. I thought about copying it over clean. I've decided not to. The wrong turns are most of what I actually learned.")
attribution("-- Wren Calder, Apprentice, College of Winterhold")

archivist_note("Archivist's Note -- closing",
    "The notebook ends here; no further volume has been located. Whether Wren Calder completed her apprenticeship, we do not know. We have chosen to preserve her crossed-out lines rather than clean them up, in keeping with her own stated wish above. -- Recordkeeper's Office"
)

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "A Notebook on Atronachs and Their False Kin.docx")
doc.save(out_path)
print("SAVED:", out_path)
