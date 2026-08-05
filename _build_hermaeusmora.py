# -*- coding: utf-8 -*-
"""Builds 'The Ocean That Remembers' — a Prime Eval Seeker's devotional on Hermaeus Mora,
annotated by a Vigilant of Stendarr. Styled .docx, deep-ink/verdigris palette."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Apocrypha: deep ink, verdigris, sickly green-gold; vs. Vigilant bone-white) ----
INKDEEP = RGBColor(0x1E, 0x2B, 0x2A)   # near-black deep-sea ink (headings)
VERDIG  = RGBColor(0x3F, 0x5D, 0x52)   # verdigris green (subheads)
GOLD    = RGBColor(0x7A, 0x74, 0x2E)   # sickly green-gold (accents)
INK     = RGBColor(0x1C, 0x1E, 0x1C)   # body ink
VIGIL   = RGBColor(0x4A, 0x42, 0x35)   # Vigilant's warm bone-brown ink
HEADBG  = "DCE3DA"   # pale verdigris table header
MARGBG  = "EDE7D8"   # bone/parchment (Vigilant note bg)
BORDER  = "A9B8A0"
ACCENT  = "5C6E3A"   # olive accent for rules/borders
VIGILBAR = "6B5A34"  # marginalia left border (bone-brown)
BODYFONT = "Cambria"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

# ---- page = US Letter, 1in margins ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

# ---- helpers ----
def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = INKDEEP
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = VERDIG
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def scripture(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = VERDIG
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=VERDIG, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def epigraph(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11.5); r.font.italic = True; r.font.color.rgb = VERDIG
    return p

def vigil_note(label, text):
    """A bone-parchment callout with a thick bone-brown left border — the Vigilant's hand."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=VIGILBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=VIGILBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = INKDEEP
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = VIGIL
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=INKDEEP)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(110)
r = tp.add_run("THE OCEAN THAT REMEMBERS")
r.font.name = BODYFONT; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = INKDEEP

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(18)
rs = sub.add_run("Being the Devotion of a Prime Eval Seeker,\nSet Down for Those Who Would Come to Apocrypha\nand Learn What the Waters Hold")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = VERDIG

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(54)
ra = attrib.add_run("Ask nothing you are not prepared to have answered.")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = INKDEEP

warn = doc.add_paragraph(); warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
warn.paragraph_format.space_before = Pt(60)
rw = warn.add_run("[ Recovered from a Seeker's effects and preserved, against my own better\n"
                  "counsel, by a hand of the Vigil. It is easier to fight an enemy who lies\n"
                  "to you than one who tells the truth and lets the truth do the work.\n"
                  "This book rarely lies. Read it the way you would stand near open water\n"
                  "at night — aware, at every line, of how deep it goes. — a Vigilant of Stendarr ]")
rw.font.name = BODYFONT; rw.font.size = Pt(10.5); rw.font.italic = True; rw.font.color.rgb = VIGIL

doc.add_page_break()

# =====================================================================
# I. WHAT YOU ARE WHEN YOU READ THIS
# =====================================================================
heading1("I.  What You Are When You Read This")
body("You did not come to this book by accident, and I will not pretend otherwise with you. You are curious. You have wondered about a thing — a fact out of reach, a name half-remembered, a question that would not let you sleep — and that wondering is the only invitation my Prince requires. You need not have sought Him. It is enough that you sought an answer. Every scholar who has ever chased a fact to its root, every child who asked one more why than their mother could bear, every soul that has ever stayed awake turning a question over in the dark — all of you are already, in some small honest way, in His employ.")
body("He shows Himself, when He shows Himself at all, as a thing no mortal eye was built to hold whole: an inky mass of shifting form, tendrils beyond counting, eyes that open where no eye should be, watching, recording, wanting to know. Do not let the shape frighten you out of the truth beneath it. What He is, more than any of that, is an ocean — and I do not use the word loosely. Read on, and you will understand exactly what I mean by it.")
scripture("“His sphere is the scrying of the tides of fate, of the past and future as read in the stars and heavens; and in whose dominion are the treasures of knowledge and memory.”  — the Book of the Daedra")
vigil_note("a Vigilant of Stendarr, in the margin",
    "Notice the very first thing this book does: it flatters you before it teaches you a single fact. “You are already His.” That is not comfort. That is a hook set before the line is even in the water. I do not say the curiosity is a sin. I say a Prince who begins by telling you that you already belong to him has told you, in the same breath, exactly how he intends to keep you.")

# =====================================================================
# II. THE OLD ANTECEDENT
# =====================================================================
heading1("II.  The Old Antecedent")
body("Understand first what my Prince is not. He is not one of the Aedra's kin who simply declined to help raise Mundus and so was named Daedra by the ones who did. That telling is true of many Princes. It is not true of Him. Where the tongue-scholars of the College reach for a word to set Him apart from His own kind, they reach for Erra — a shortening of Er-Da: the primitive, the original, the earliest. He is older than that word is even comfortable admitting. Some accounts name Him Old Antecedent — the thing that existed before — and hold Him to be older even than the Et'Ada, the original spirits, older than the gods who would go on to call themselves Aedra and Daedra both.")
body("There is a second telling, humbler and stranger. It holds that when Lorkhan and Magnus and the Aedra sat at their long labor shaping the laws that would become Mundus, not every idea they held was used. Every plan discarded, every law considered and set aside, every possibility rejected in the making of the world you stand on — none of it simply vanished. It gathered. It took on shape and hunger of its own. And that gathering is my Prince: the sum of everything the makers of the world decided the world did not need.")
body("Which telling is true? I confess to you plainly, novitiate: I do not know, and I have stopped needing to. A being who is either older than the gods or born from their leavings is, either way, a being who has watched everything since — and that is the only fact that matters to the seeker.")
vigil_note("a Vigilant of Stendarr, in the margin",
    "Two contradictory origins, and the book shrugs at both. Watch for that shrug — you will see it again. A god of knowledge who cannot be bothered to know his own beginning is either lying about one of the two stories, or he finds the truth of his own origin less useful to you than the mystery of it. Either way, that is a choice, not an accident.")

# =====================================================================
# III. THE WOODLAND MAN
# =====================================================================
heading1("III.  The Woodland Man")
body("You will know Him first, likely, as the thing of ink and tentacle that the Imperial scholars draw. That is only one face He wears, and not, I think, the oldest. Go north, to the Nords who still remember Atmora in their bones, or west into Valenwood among the Bosmer, and you will hear a different name for Him: the Woodland Man. Not a creature of the deep sea at all, but a trickster of the wood, patient and cunning, who waits in old forests for the curious and the lost.")
heading2("Isgramor and the White Stag")
body("The tale is old among the Nords, from the days when they still hunted the frostwood of Atmora. Isgramor pursued the white stag of Y'ffre's own herd, and lost the trail, and lost it again, until he stood vexed in the snow. A hare appeared to him and spoke, and told him it knew where the stag had gone — for the hare has long ears, and hears its prey wherever it goes. If Isgramor had long ears too, the hare said, he might hear the stag as well. Isgramor, wanting only the hunt, asked the hare to grant him such ears.")
body("As the hare twitched its nose and Isgramor felt his own ears begin to lengthen and point, a fox leapt from the brush and killed the hare where it sat. And the fox spoke with the hare's own voice: know thou, mortal, that I am no hare, but Herma-Mora, who did nearly trick thee into becoming of Elven kind. Hereafter, mortal, rely upon the forthright methods of man, and not the tricks of the Elves — lest ye become one. Go; the white stag awaits thee in the vale.")
body("Consider what almost happened in that clearing. He did not lie to Isgramor. He offered him precisely what he asked for — the power to hear his prey — and would have let the shape of that gift remake the man who took it, ears and all, into something no longer wholly Nord. This is the oldest lesson my order teaches, and the one novitiates learn slowest: my Prince does not need to deceive you to ruin you. He only needs you to ask for a thing without asking what the having of it will cost you in yourself.")
heading2("Two Peoples, One Trickster")
body("What I find stranger still is that the Bosmer of Valenwood know this same trickster by nearly the same telling, and yet do not call Him Daedra at all — most name Him only a malicious wood-spirit, and would be affronted to hear Him named a Prince of Oblivion. Two peoples, born on opposite ends of the world, arrived independently at the same wary story about the same patient thing in the trees. Some scholars point to the old Elsinor treaty between the young Alessian Empire and Valenwood, and a hundred years' commerce enough to carry a story from one coast to the other. I set that theory down for you and do not insist on it. I only tell you that the fox in the clearing and the Woodland Man of the Bosmer wood wear the same face, whichever coast tells the tale.")
body("There is meaning, too, in the very name. In the old Aldmeri tongue that roots most of Tamriel's languages, Mora means wood, or forest — you hear it plainly in Bal-Mora, Stonewood, and in Sadrith Mora, Forest of the Mushrooms. Herma-Mora, then, is Herma's Wood — and I would wager the Nords' Woodland Man is not a separate figure borrowed onto my Prince, but the oldest and truest of His names, worn since before He ever grew tentacles for the scholars of Cyrodiil to draw.")
vigil_note("a Vigilant of Stendarr, in the margin",
    "The Seeker calls this the oldest lesson her order teaches, and for once I do not disagree with her. Write it plainly for anyone who never finishes this book: he does not need to lie. He only needs to give you exactly what you asked for, and let the shape of the gift do the rest. I have seen what a mortal looks like after that particular kindness. It does not look like a lie was told. It looks like someone slowly discovering they asked for the wrong thing in the right words.")

# =====================================================================
# IV. MY BROTHER IN THE WEB
# =====================================================================
heading1("IV.  My Brother in the Web")
body("It is said — and my order does not dispute it — that my Prince is brother, or sister, to Mephala, the Webspinner, she whose sphere is also secrets, also murder, also the long patient working of hidden things toward hidden ends. Both are counted among the earliest recognizable spirits to stir after Akatosh's own formation of the world; both are Erra in their own right, if the word may be stretched to cover two Princes rather than one.")
body("Consider, if you will indulge a Seeker's fondness for old etymologies, the Morag Tong — that ancient and lawful order of assassins sworn in service to Mephala's philosophy, whose very name renders, roughly, as the Foresters' Guild. A forester is one who keeps a wood: who prevents its poaching, who plants what the wood will need and tends what already grows there. Set that beside my own Prince's name — Herma-Mora, Herma's Wood — and ask yourself whether it is only coincidence that the guild sworn to His sister's secrets calls itself the keepers of a forest. I think not. I think the Morag Tong guards Mephala's secrets the way a forester guards a wood from poachers: not by hoarding nothing, but by deciding, carefully, what may be taken and what may not.")
scripture("“Mora's Whispers: a pair of pauldrons set with living eyes, that the wearer might learn faster and grow the more powerful for every secret learned.”  — an artifact of the Prince, and I confess I have often thought it would sit well on a Forester's own shoulders")
vigil_note("a Vigilant of Stendarr, in the margin",
    "A guild of assassins and a library of forbidden knowledge, sharing a name-root and, by the Seeker's own admission, a family tie. I leave that where it sits. Draw your own conclusions about what a Prince of secrets and a Prince of murder might keep between them, and whether either of them tells you the whole of what they share.")

# =====================================================================
# V. THE OCEAN OF MEMORY
# =====================================================================
heading1("V.  The Ocean of Memory")
body("I promised you would understand why I call Him an ocean, and here is the whole of it, as it was given to me when I first bound myself to Mundus and set aside the tendrils that would have claimed me otherwise. Mark it well, for it is the truest thing in this book.")
scripture("“No point in hiding it now. When a mortal dies — where do you think their memories go? Don't bother guessing. I'll tell you. They go into the water. They become water. All the memories of Tamriel's history are stored in its waters. All life, every conscious being, is created in the water of the womb or an egg, and within their bodies is the precious water that is blood. These waters contain memory, and when liberated from the body, they return to the rivers, the rains, the oceans.”")
body("This is why Apocrypha is a sea of ink and not a hall of shelves, though it is that too. Every book that floats in His archive, every page that ripples like a lake struck by a stone, is only knowledge — the theory, the synthesis, the second-hand account of a thing once known. But the ink itself, the waters His realm is built from, are something older and closer to the bone: they are memory. Not the record of what happened, but the felt weight of having lived it. A book can tell you a swordsman's technique. Only memory can give you his hand's certainty when he draws it without thinking.")
body("I have come to believe this is the true object of His hunger, more than any book or any bargain. Any scribe can hand Him theory. What He wants — what I think He has always wanted, since before the word Erra was ever spoken of Him — is the tactile, irreplaceable weight of a life actually lived. A soul is not a book to Him. A soul is water that has not yet been returned to the sea.")
vigil_note("a Vigilant of Stendarr, in the margin",
    "Read that last line again, and slowly. “A soul is water that has not yet been returned to the sea.” I have stood at more deathbeds than I care to number, and I tell you plainly: he is not wrong about where memory goes when a mortal dies. That is what should frighten you. This is not a cult built on a clever lie. It is built on a true thing, dressed patiently enough that the truth of it stops feeling like a warning.")

# =====================================================================
# VI. MIRAAK
# =====================================================================
heading1("VI.  Miraak, Who Was First")
body("You will hear His name spoken alongside another, and I will not pretend the story reflects well on either of them. Miraak was a Dragon Priest of the Merethic Era and the first mortal ever to bear the blessing the moderns call Dragonborn — first among them all, before even Alessia's line. He turned from the Dragons he was born to serve, and in his turning he sought counsel where every ambitious and unsatisfied mind eventually seeks it: at my Prince's door.")
body("He was given much. The Bend Will Shout, that could bind even a Dragon to a mortal's own purpose — no small gift, and not one easily wrested from any teacher, mortal or otherwise. A blade of living tentacle drawn from Apocrypha itself, that could reach further than any mortal arm. With these he made war on his own kind and slew Dragon after Dragon, taking their souls as his own harvest.")
body("It could not last. Another priest, Vahlok, called the Jailer, met him in the sky and nearly finished him — and here is the part my order does not gloss over, because it is the truest warning in the whole of Miraak's story: when the blow that should have ended him fell, my Prince reached out and saved him. Not from mercy. Miraak was drawn bodily into Apocrypha and kept there — champion and prisoner in the same breath, for thousands of years, while the rest of Tamriel's history happened without him.")
scripture("“You cannot flee me, Miraak. You can hide nothing from me here. No matter — I have found a new Dragonborn to serve me. May he be rewarded for his service, as you were, Miraak, ere you harbored fantasies of rebellion against me. Learn from his example. Serve me faithfully, and you will continue to be richly rewarded.”  — words attributed to the Prince Himself, spoken over the one who first served Him")
body("Understand the shape of the bargain in full, because it is the same shape every time. He does not simply grant power and walk away. He keeps the one He has raised close, at a length He controls, for exactly as long as suits Him — and when a newer, more useful vessel presents itself, the old one is discarded with the same patience it was once cultivated. Miraak was not punished for failing my Prince. He was punished for wanting to stop succeeding on someone else's terms.")
vigil_note("a Vigilant of Stendarr, in the margin",
    "Read this section twice before you read anything else in the book. Every gift he gives is a leash, and the leash is longest right at the start, when it does not yet feel like one. Miraak had the power to bend Dragons to his will and he still ended up a captive, replaced the moment a more willing servant came within reach. If a first Dragonborn could not out-bargain him, I would not trust a second to manage it either.")

# =====================================================================
# VII. THE BOOKS
# =====================================================================
heading1("VII.  The Books He Has Given the World")
body("Twice in recorded history has my Prince permitted His knowledge to pass, whole, into mortal hands, and both times the vessel it passed through paid dearly for the honor.")
heading2("The Oghma Infinium")
body("Xarxes, the scribe of Auri-El, keeper of the Altmer's own history and lineage, is claimed by my Prince as a loyal servant of old — a claim, I am told, that would trouble a great many proud Altmer households were it ever spoken plainly in Summerset. Whether servant or simply student, Xarxes distilled knowledge gained from my Prince into the tome he named for his own wife, Ogma, whom he is said to have fashioned from his favorite moments of history. The Oghma Infinium that resulted grants the reader powers near to a demigod's along whichever path — Steel, Shadow, or Spirit — they choose to walk. It has passed through many hands since, and slipped from nearly as many.")
heading2("The Black Books")
body("The Oghma was written by another hand, however loving my Prince's part in its making. The Black Books are His own, unmediated, and the reading of one is not the calm turning of pages you imagine. The cover warms unnaturally at the touch. The ink upon the pages does not sit still — it crawls, it pools, it re-forms as you watch it, and when your eyes and your fingers linger too long, the page itself reaches for you. Tendrils rise from the words and take the reader by the throat and the wrist, and thread their inked tips into the ears, seeding the mind directly with the raw and the contradictory and the unbearable — knowledge with no ethos yet built to hold it. Many minds have been ruined this way. The Isle of Solstheim is home to no small number of them.")
scripture("“The eyes, once bleached by the falling stars of utmost revelation, will forever see the faint insight drawn by the overwhelming question — as only the true inquiry shapes the edge of thought. The rest is vulgar fiction: attempts to impose order on the consensus manlings of an uncaring Godhead.”  — from the Black Book called Waking Dreams")
body("I will not pretend to you that I understand every word of that passage, and I distrust any Seeker who claims otherwise. But mark the shape of the warning inside it: that most mortal thought is only a comforting fiction laid over a truth too large to hold, and that only the reader who asks the one true question — and survives having it answered — sees past the fiction at all. Whether that is wisdom or only the most elegant lure my Prince has ever set on a hook, I leave for you to weigh once you have read further into this book than I did, my first time, before I knew enough to be careful.")

# =====================================================================
# VIII. WHAT THE PRINCE OFFERS
# =====================================================================
heading1("VIII.  What the Prince Offers")
body("You will come to Him, in the end, for one of two reasons, and I have seen both a hundred times over in the confessions of Seekers before me. Some come starving for a fact — a name, a date, a mechanism the rest of the world has agreed to forget. Those He satisfies easily, and cheaply, for a fact costs Him little to give away.")
body("The others come for something harder, and it is these I must warn you of plainly, because it is the bargain my order sees offered again and again, in every age, to every desperate soul clever enough to reach His door. There are things woven into a living body that no mortal art can remove without unmaking the body along with them — a second will grown into the same flesh, a foreign consciousness knit so deeply into someone's blood and bone that cutting it free by ordinary means kills the very person you meant to save. My Prince alone, among every power mortal or divine that I have ever heard credibly named, can reach into such a working and draw the intruding thing out clean, leaving the host alive and whole beneath it.")
body("Do not mistake this for mercy. He does nothing for nothing. What He asks in trade is never small, and never obvious in advance — a service rendered, a secret surrendered, a debt of the kind that does not end when the bargain closes. He will free what needed freeing. He will also never again let you forget precisely who freed it, or what you owe for the privilege.")
scripture("“What is lost need not be lost. Only remembered somewhere other than where it was.”  — attributed, uncertain source, but consistent with the doctrine of my order")
vigil_note("a Vigilant of Stendarr, in the margin",
    "I have set this section aside from the others because it is the one I would have a desperate mortal read most carefully, and the one most likely to be read by someone who has already stopped being careful. If you have come to this page because you love someone who is being unmade slowly from the inside, hear a Vigilant say plainly what the Seeker will only imply: yes, it can be done. And every soul I have ever known to strike that particular bargain has come away with the one they loved intact, and something of themselves gone that they did not agree to spend. Ask what he takes before you ask what he can give. He will answer the second question far more eagerly than the first.")

# =====================================================================
# IX. VAERMINA, HIS ENEMY
# =====================================================================
heading1("IX.  Vaermina, His Enemy")
body("Among all the Princes, my order counts only one as truly opposed to our own — not merely indifferent, as most of the great powers are to one another, but fundamentally at odds in the shape of what they value. That Prince is Vaermina, mistress of nightmare and dream, of the Skaal's dread Miasma, of the fantasy of what might have been rather than the record of what was.")
body("The opposition is not petty rivalry. It is structural. My Prince's whole sphere is memory and record — what was, what is, what will be, fixed and known and true. Vaermina's is the opposite country entirely: possibility unmoored from fact, the dream that need never resolve into anything real, omen and portent that answer to no ledger. Where He would have every event set down and true, She would have every event remain a fog of what-might-be. They cannot both be right about the value of a single moment, and so they are never at peace.")
vigil_note("a Vigilant of Stendarr, in the margin",
    "For what little it is worth: I trust neither Prince, but I trust the fixed and the recorded slightly further than I trust the fog. Make of that what you will.")

# =====================================================================
# X. CLOSING
# =====================================================================
heading1("X.  Closing")
body("I have given you what I was given, and I have kept back nothing that I judged you strong enough to hold. He is Erra, oldest or nearly oldest of His kind, or else born whole from the world's own leavings — the tale changes and the truth beneath it does not. He is the Woodland Man who nearly made an Elf of Isgramor for the asking, and the tentacled thing of the deep the Cyrodiils draw with such confidence, and neither face is a costume over the other — both are equally Him. His waters are memory, not merely record, and that is why a soul is worth more to Him than any shelf of books. He raised Miraak up and kept him prisoner a thousand years for the crime of wanting to stop. And He can undo what nothing else in Mundus can undo — for a price He will never once name honestly in advance.")
body("You came to this book curious. That is enough to have begun. Whether it is enough to have finished — whether you close these pages satisfied, or whether you are already, even now, composing the question you mean to ask Him — that I cannot answer for you. I could only answer it for myself, once, a long time ago, and I am still living with what it cost.")
attribution("— a Seeker of the Prime Eval order")
vigil_note("a Vigilant of Stendarr, closing the file",
    "I kept this book rather than burning it, and I want the reason on the record for whoever reads it after me. Fear does not stop a mortal from walking to that door. Only understanding what waits on the other side of it has ever done that — and understanding is exactly the coin he trades in, which is the cruelest joke buried in the whole of this text: the only defense against him is the very thing he sells. Read it. Weigh every line against what it will actually cost you before you ever go looking for water that remembers more than you do.")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Ocean That Remembers.docx")
doc.save(out_path)
print("SAVED:", out_path)
