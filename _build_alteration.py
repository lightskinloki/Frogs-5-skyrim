# -*- coding: utf-8 -*-
"""Builds 'On the Negotiable Law' -- a Tier-4 school-of-magic treatise on
Alteration. Single authorial voice throughout (Cassian Threll, Master of
Alteration, College of Winterhold -- living, teaching, confident), matching
the house pattern established by 'On the Doctrine of Essences' and 'Soul
Trap Tome': one professor's argued, opinionated, personally-honest treatise,
no running second voice, no narrative plot. Brief, calm archivist framing at
open and close only, matching 'On the Structure of Magical Law'. Cross-
checked against that text for somatic-component consistency (Alteration's
natural affinities: ANCHOR, SHIFT, BIND, STABILIZE) and against verified
lore (Bravil / Te'o Bravilius Tasus; Reality & Other Falsehoods; the
Earthbones cosmology). Styled .docx, slate/silver palette.
Run: python _build_alteration.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Alteration: slate/silver-blue) ----
SLATE   = RGBColor(0x2C, 0x3E, 0x4A)
STEEL   = RGBColor(0x41, 0x5A, 0x68)
SILVER  = RGBColor(0x5E, 0x74, 0x7E)
INK     = RGBColor(0x1C, 0x1C, 0x1E)
AMBER   = RGBColor(0x6B, 0x4A, 0x1E)
HEADBG  = "D9E1E4"
MARGBG  = "EFE6D8"
BORDER  = "9FB1B8"
ACCENT  = "4A6472"
AMBERBAR = "6B4A1E"
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = SLATE
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = STEEL
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = SILVER
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=SILVER, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def archivist_note(label, text):
    """A brief, calm callout -- used only at open and close. Not a running voice."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=AMBERBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=AMBERBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = AMBER
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = AMBER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=SLATE)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(100)
r = tp.add_run("ON THE NEGOTIABLE LAW")
r.font.name = BODYFONT; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = SLATE

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(16)
rs = sub.add_run("Being a Treatise on the True Nature\nof the School of Alteration")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = STEEL

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("Cassian Threll\nMaster of Alteration, College of Winterhold\n4E 201")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.italic = True; ra.font.color.rgb = SLATE

doc.add_page_break()

archivist_note("Archivist's Note -- College of Winterhold",
    "This treatise is assigned to Adept-standing students of Alteration and above. Master Threll continues to teach the introductory sequence each term; students with questions about the arguments below are encouraged to raise them with him directly during office hours, which he keeps faithfully. -- Recordkeeper's Office"
)

# =====================================================================
# I. THE FALSE PARTITION
# =====================================================================
heading1("I.  The False Partition")
body("Every apprentice is taught the same comforting lie in their first week at any college of magic: that the six schools name six different kinds of power. Destruction burns. Restoration mends. Illusion deceives. Conjuration summons. Alteration merely adjusts. The lie is comforting because it is orderly, and orderly things are easy to grade on an examination.")
body("It is also false. There are not six powers. There is one power, and it does not originate in the caster at all. Magicka descends from Aetherius, pouring down through the rifts Magnus tore in the sky when he and the greater part of the et'Ada fled the unfinished work of Mundus rather than pay its price in full. Every mage who has ever lived draws from that same falling light, whether they call it down to burn a bandit's flesh or to knit a child's broken arm. The schools are not six rivers. They are six words for the same river, spoken by scribes who needed a filing system more than they needed the truth.")
body("I do not say this to diminish my colleagues in Destruction or Restoration. I say it because the filing system has a cost that no one troubles to mention in the first week, or the fifth year, or ever: it has taught every mage in every college on this continent to stop asking what a spell actually does to reality, so long as they can say which drawer it belongs in.")
body("I have taught Alteration at this College for the better part of two decades. I will tell you plainly, as I tell every student who reaches this level: I believe Alteration is the one school honest enough to keep asking the question the filing system was built to end. That is the argument this treatise makes. Read it as an argument, not a settled fact -- I have colleagues in this very College who dispute parts of it, and I would rather you knew that than mistake my confidence for consensus.")

# =====================================================================
# II. THE SIEGE OF BRAVIL
# =====================================================================
heading1("II.  The Siege of Bravil")
body("Let me offer the proof before the philosophy, because philosophy unearned by evidence is only opinion dressed for company.")
body("When the slave-queen Alessia's rebellion had broken the Ayleid dominion across most of Cyrodiil, the city that would come to be called Bravil, on the Niben, held against her for far longer than its garrison should have allowed. Four times her Centurions -- the field-officer rank of her own young army, no kin to the Dwemer's mechanical namesakes -- took the streets by open combat, drove the defenders back within their own walls, and camped their soldiers to hold the gates through the night. Four times, dawn found every occupying soldier dead and the Ayleids walking their own battlements again, and no one among the living had heard a fight.")
body("The easy explanation -- that the elves had never truly lost the city, that some trick of theater had convinced weary soldiers they had won a battle they had not -- satisfied no one who had actually stood in those streets and watched the enemy break. The harder truth took one Centurion's unusual patience to find: Te'o Bravilius Tasus, who noted shallow handholds worn into the stone far above any soldier's reach, and, once, a single footprint pressed into riverbank mud beside a garrison that had posted no one near the water that night. The Ayleids were not holding the city by force of arms at all. They were holding it by making the stone climbable to hands it had no business admitting, and the river breathable to lungs that should have drowned in it, until an army convinced of its own victory had nothing left to fight. The city keeps his name still, worn down by the centuries the way his own find wore down the stone: Bravilius become Bravil.")
quote("\"They found us four times and lost us four times, for what they were hunting for was men, and men were not what waited above them in the wall.\"  -- fragment, Ayleid garrison marker, provenance disputed")
body("Here is what the College will not put on an examination, because it unravels the examination's whole premise: no lockpick, ward, or fire-spell ended that siege. Reality itself was renegotiated -- stone made climbable, water made breathable, the defenders' own shapes made other than what they were -- and a superior army could find nothing to fight because there was, by any honest accounting of what is normally true, nothing there to find. This is not a parlor trick with a generous name. This is the oldest recorded use of this school, and it very nearly held an empire at bay by itself.")

# =====================================================================
# III. THE LIE OF UTILITY
# =====================================================================
heading1("III.  The Lie of Utility")
body("Ask any journeyman what Alteration is for and you will hear the same short list, recited like a nursery rhyme: hardened skin for a fighter who forgot their armor, a lighter load for a packhorse, breath underwater for a diver, an open lock for a thief too proud to carry picks. Useful. Modest. A trade for people who did not have the stomach for real magic.")
body("I have taught this list myself, to students who wanted nothing more from it, and I do not regret teaching it to them. But I want to set down plainly what that list is actually made of, because no one before me has troubled to say it in a single sentence: every spell on it is a small, deliberately unthreatening demonstration that the properties of matter are not fixed. Flesh becomes as resistant as oak, and then as stone, and then, at the furthest reach any living mage has documented, as resistant as the black metal drawn from fallen stars. Lungs that require air stop requiring it. A body's own weight -- weight, which every apprentice is taught is the most basic fact an object can have -- becomes a number the caster may simply choose.")
body("Say it plainly and the modesty falls away at once. This is not a trade for the timid. This is the school that has quietly demonstrated, in every college on this continent, for as long as there have been colleges, that the physical world holds no fact so basic the right working cannot set it aside.")

# =====================================================================
# IV. WHAT LAW ACTUALLY IS
# =====================================================================
heading1("IV.  What Law Actually Is")
body("Master Imedril of Artaeum, five centuries dead and still the clearest mind our discipline has produced, wrote the sentence that every mage in this college now recites without understanding it: Magicka does not respond to desire. It responds to relationship. A spell is not an imposition on reality. It is a negotiation with it.")
body("I ask my colleagues in the other schools to sit with that sentence honestly for a moment, because I do not believe most of them have. A Destruction mage who casts Firebolt is negotiating too, by Imedril's own account -- COMMAND and PROJECT, forcing fire to manifest and travel. But COMMAND is the loudest, crudest term in that negotiation, the one that most resembles simply shouting. It is negotiation the way a bandit's demand at knifepoint is negotiation.")
body("Look instead at where our own discipline's fluency actually lives, by the Archivists' own accounting: ANCHOR, SHIFT, BIND, STABILIZE. Not force. Rootedness. Adaptation. Persistence. Correctness. These are not the vocabulary of a school that pushes against law from the outside. They are the vocabulary of a school that has gone inside law's own house and learned to speak its language on its own terms -- to root oneself so thoroughly in a claim about reality that reality finds it easier to agree than to argue, to insist that a changed state is not a temporary exception but the new and simple truth, until the world stops bothering to disagree.")
lore_table(
    ["Component (per Imedril)", "What it says, in my reading"],
    [["ANCHOR", "I am not asking the world to change around me. I am becoming, myself, the fixed point the world must now account for."],
     ["SHIFT", "I do not command the law where to go. I have made myself fluent enough in its own inclination that it goes where I already am."],
     ["BIND", "This is no longer a request. It is a fact now, and facts do not require the world's permission to persist."],
     ["STABILIZE", "I am not holding this against reality. I have convinced reality that this was always the true state, and it is arguing with itself, not with me."]],
    [2.1, 4.4])
body("This is why Alteration alone, of the six schools, keeps producing practitioners who cannot stop asking what law itself is made of. The other schools borrow law's power and spend it. Ours learns law's grammar and speaks it back. A mage who has genuinely done this for long enough stops experiencing the physical world as fixed and begins experiencing it as merely well-defended -- persuasive, consistent, extremely convincing, and, underneath all of that, negotiable.")

# =====================================================================
# V. ILLUSION'S LESSER TRUTH
# =====================================================================
heading1("V.  Illusion's Lesser Truth")
body("My colleagues in Illusion will object to everything in this chapter, and I invite the objection, because the distinction is worth stating precisely rather than gesturing at.")
body("An Illusion spell convinces a mind that something is true. A Fear spell does not summon a real threat; it manufactures, inside a single skull, the conviction that a threat is present, and the body obeys a conviction exactly as readily as it obeys a fact. This is powerful, and I do not belittle it. But it is also, by its own definition, a lie that only one mind believes. Break the caster's concentration, and the target's own reasoning may claw its way back to what was actually true the entire time, because the entire time, something was actually true, and the Illusion mage never touched it.")
body("An Alteration spell that hides a body among the stones of a wall, or grants it breath in water where breath is not naturally owed, produces no such gap between belief and fact. There is nothing for a clear-headed target to claw their way back to, because nothing was ever merely believed. The stone really did become passable. The lung really did stop requiring air. Any observer, however skeptical, however unaffected by any spell of persuasion, will find the same altered fact waiting for them that the target finds. This is not illusion's cousin. This is illusion's superior, doing the harder work honestly instead of the easier work by suggestion.")
quote("A lie only one mind believes can be argued out of that mind. A truth every mind must now agree to has no mind left to argue with.")
body("I will grant my colleague Master Delvarn in Illusion the point he has made to me more than once, in more heated terms than I am reproducing here: usefulness is not the same axis as honesty, and a mage in the field rarely has the luxury of choosing the philosophically superior tool over the one that actually gets them through the door. He is right about that. I maintain the distinction anyway, because knowing which tool you are actually holding matters even when you would have reached for it either way.")

# =====================================================================
# VI. THE TREMBLING BONES OF THE WORLD
# =====================================================================
heading1("VI.  The Trembling Bones of the World")
body("We are taught, correctly, that the greater et'Ada who remained after Mundus was made poured themselves into its foundations rather than abandon the unfinished work entirely. Akatosh gave time its shape and its single direction. Kynareth gave the elements their behavior and the sky its wind. Every law a mage of any school negotiates with is, at bottom, a fragment of one of these old and willing sacrifices, still holding a shape it agreed to hold before any mortal existed to ask it to.")
body("Consider, then, what it actually means when a mage roots their own flesh in the resilience of stone, or teaches their own lungs to disregard the law that Kynareth's own gift established for breathing creatures, or fixes a changed state so thoroughly that the world's own tendency toward its original arrangement -- what my colleagues in Destruction call, correctly, entropy, and what I prefer to call law's memory of what it used to be -- simply gives up trying to reassert itself. We are not borrowing a convenient exception. We are placing a small, specific, unmistakable pressure directly against a promise a god made before recorded history, and finding, every time, that the promise yields.")
body("It always yields to a small enough pressure. That is the comfort every teacher of this school offers a nervous apprentice, and it is true as far as it goes. I no longer find it as comforting as I once did. A promise that yields to a small pressure has not been shown to be unbreakable. It has only been shown to have never yet been pushed hard enough to learn where its edge is.")

# =====================================================================
# VII. THE APEX OF THE DISCIPLINE
# =====================================================================
heading1("VII.  The Apex of the Discipline")
body("I will state the final claim of this treatise as plainly as I am able, because plainness is the only honesty available to a claim this large. If law itself is a promise rather than a wall -- if Akatosh's time and Kynareth's elements and every other foundation are agreements the world's makers chose to keep rather than facts the world was simply born already possessing -- then a mage who has become fluent enough in the grammar those makers used has not learned a trick. They have learned the same language the makers themselves were speaking when they built the promise in the first place.")
body("I do not claim to have reached that fluency, and I want to be direct about why I am telling you this rather than leaving it as an implication for you to draw yourself. Every serious student of this school eventually asks their instructor, in some form, whether the old warnings about Alteration are exaggerated. I do not think they are. The College's own record of Ontological Detachment among overreaching practitioners is not folklore; it is filed, dated, and available to any Adept who requests it. I have read the file. I teach this school anyway, because I believe the honest study of a real danger is safer than the superstitious avoidance of it, and because I think a mage who understands exactly what the edge looks like is less likely to wander over it by accident than one who has only ever been told, vaguely, not to look.")
body("I will not pretend to know what waits at the far end of that fluency, fully realized. I know of no living mage who has reached it and returned able to describe it in terms the rest of us could use. I raise the question because a treatise that stops short of its own argument's real implications has not finished its work, whatever comfort the stopping short might offer both of us.")

# =====================================================================
# CLOSING REMARKS
# =====================================================================
heading1("Closing Remarks")
body("I have been teaching this material long enough to know which parts of it students remember and which parts they discard the moment the examination ends. What I hope survives, more than any single claim in this treatise, is the habit underneath all of them: ask what a spell is actually doing, not merely what drawer it has been filed into. That habit is the whole of what I have to teach, dressed up, across seven chapters, in the specific clothing of my own discipline.")
body("I remain at this College, teaching this course, most terms of most years. Students with objections to anything written here are welcome to bring them to me directly. I have changed my mind about smaller points in this treatise more than once, and I do not expect this to be the last revision it receives.")
attribution("-- Cassian Threll, Master of Alteration, College of Winterhold, 4E 201")

archivist_note("Archivist's Note -- closing",
    "This text has been reviewed by the College's Alteration faculty and is consistent with current teaching practice, though Master Threll's Chapter VII is understood among his colleagues to represent his own position rather than settled College doctrine. We reproduce it as submitted. -- Recordkeeper's Office"
)

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "On the Negotiable Law.docx")
doc.save(out_path)
print("SAVED:", out_path)
