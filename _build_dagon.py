# -*- coding: utf-8 -*-
"""Builds 'The Red Legions -- A War College Chronicle of Mehrunes Dagon' -- a
single-voice military-doctrine text, NOT a doctrinal/cult text (that ground is
already fully covered by 'The Dawn Is Yours' / Mythic Dawn book -- this is
deliberately a different genre entirely: history and strategy, not scripture).

VOICE: an instructor at the Imperial Battlespire, the war college Dagon himself
crippled once (3E 398, in league with the usurper Jagar Tharn). Real personal
and institutional stake -- not a neutral historian. Writing generations after
the Oblivion Crisis (3E 433), compiling the full documented pattern of Dagon's
interventions as required reading for war-college cadets. Thesis, stated as
hard operational doctrine rather than philosophy: Dagon does not invade for
ground. He weaponizes an injustice that already exists. The only real defense
is addressing what he exploits, not merely fortifying against him.

Content: Dagon's mythic-era origin at Lig (framed skeptically -- the author
flags that no reliable record this old exists, but extracts the pattern
anyway); the sack of Mournhold and the fight with Almalexia and Sotha Sil,
with the nearly-simultaneous fall of the Reman line to Akaviri betrayal noted
as the author's own caveat that correlation isn't causation; the Battlespire
assault of 3E 398 (the author's own institution, own stake); the Oblivion
Crisis of 3E 433 (Mythic Dawn, the gates, the Amulet of Kings, Martin Septim).
One artifact is deliberately NOT given a firm name -- the GM's transcript
source called it "the Sword of the Moon River," which reads as a garbled
transcription of something real; per GM instruction, the text notes the name
as disputed in the author's sources rather than inventing or silently
"correcting" it. All content reconstructed in original prose, not the source
video's phrasing, cross-checked against general ES lore before drafting.
Styled .docx, martial ash/iron/dried-blood palette -- deliberately NOT the
warm ember/gold of the Mythic Dawn book or the cold glacial-blue of the
Umbra book.
Run: python _build_dagon.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (war college doctrine: ash, iron, dried blood -- NOT ember/gold, NOT glacial blue) ----
IRON     = RGBColor(0x2A, 0x28, 0x26)   # near-black iron (headings)
ASH      = RGBColor(0x4A, 0x44, 0x40)   # ash-grey (subheads)
DRIEDRED = RGBColor(0x6B, 0x2E, 0x28)   # dried-blood maroon (accent, sparing)
INK      = RGBColor(0x1D, 0x1B, 0x1A)  # near-black body ink
FADED    = RGBColor(0x5C, 0x56, 0x50)   # faded grey (citation/caveat ink)
HEADBG   = "E4DFD9"   # ash-parchment table header
BOXBG    = "EAE4DE"   # doctrine-box bg
BORDER   = "B8ADA2"
ACCENT   = "5C4A42"   # muted ash-brown accent for rules
BOXBAR   = "6B2E28"   # dried-blood bar on doctrine boxes
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = IRON
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def kicker(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.font.name = BODYFONT; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = DRIEDRED
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def caveat(text):
    """A source-honesty aside -- the author flagging shaky evidence."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = FADED
    return p

def doctrine(label, text):
    """A boxed operational-doctrine conclusion."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, BOXBG)
    _set_cell_borders(cell, color=BOXBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=BOXBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True; lr.font.color.rgb = DRIEDRED
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = IRON
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def attribution(text):
    return body(text, italic=True, size=10.5, color=FADED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(90)
r = tp.add_run("THE RED LEGIONS")
r.font.name = BODYFONT; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = IRON

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("A War College Chronicle of Mehrunes Dagon,\nCompiled for the Instruction of Cadets")
rs.font.name = BODYFONT; rs.font.size = Pt(12.5); rs.font.italic = True; rs.font.color.rgb = ASH

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("Compiled by a serving Instructor,\nthe Imperial Battlespire")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = IRON

doc.add_page_break()

# =====================================================================
heading1("Why This Book Exists")
body("You are here to learn to fight men, mostly, and the occasional beast, and once in a great while something that used to be a man and no longer is. This book concerns none of those. It concerns a Prince, and I will tell you plainly why a war college troubles itself with theology: because the single worst defeat this institution has ever suffered was inflicted by him, and because every cadet who leaves this hall without understanding how he operates is a cadet who will one day face the pattern again and call it something else.")
body("I am not a priest of any Daedra, his least of all, and I will not pretend to speak for his intentions the way his own faithful do. I have read their scripture. It is not this book. This book is a chronicle of what he has actually done, four times that the record allows us to trust, and what those four occasions have in common, because the common thread is the only part of him that is any use to a soldier.")
doctrine("The thesis, stated once, plainly",
    "Mehrunes Dagon does not invade for ground. He arrives where an injustice was already load-bearing, and he puts his weight on it. Every case in this book follows that shape. The defense that matters is not a wall. It is not letting the injustice accumulate in the first place.")

# =====================================================================
kicker("The first case (unverifiable, and I say so)")
heading1("Lig, and the Making of a Prince Nobody Meant to Make")
body("I will treat this one differently than the rest, because I have no confidence in it as history and every confidence in it as instruction, and a good instructor tells you which is which.")
body("The account, such as it survives, holds that Dagon did not always exist as we know him -- that in an age before Mundus had settled into the shape it now keeps, a race of tyrant rulers held an entire world called Lig in absolute bondage, and that certain lesser spirits attending the god of magic, seeing the misery of it, resolved to build a champion for the oppressed. They meant to make a Prince of pure hope. What woke instead called itself Mehrunes Dagon, and Lig did not survive his liberation of it.")
caveat("I want to be honest about the evidentiary weight of this account before a cadet mistakes it for settled fact, as I very nearly did in my own training. No living scholar has walked Lig, no surviving record predates the telling by many ages, and the tale reads, more than a little, like a story a Prince himself might enjoy having told about him. I include it anyway, because whether or not it is literally true, it teaches the lesson perfectly: hope, given no outlet but violence, becomes violence. Treat the history as suspect and the lesson as sound. That is the correct posture for most of what the Daedra tell us about their own origins.")
body("What the account gets right, whether or not the particulars are, is the mechanism I will spend the rest of this book demonstrating: he did not arrive as an occupying army looking for territory. He arrived because there was already a population with no more patience for its own chains, and he became the shape their patience ran out into.")

# =====================================================================
kicker("The second case")
heading1("Mournhold, and What a Grudge Can Summon")
body("Better documented, because mortals were doing the summoning this time, and mortals keep better records of their own schemes than Daedra keep of theirs.")
body("Two petitioners, each nursing an old grievance against the Duke of Mournhold, made the offering between them and called Mehrunes Dagon down on the city with their Duke's own blood as the price of admission. What followed was not a battle. It was closer to a demonstration. The upper city burned in a single night, the temple district was torn from its footings, and the Prince's lesser servants harried the population out of hiding and into the open where an army waited that owed obedience to no living general.")
body("Two of Morrowind's own great powers, Almalexia and Sotha Sil, fought him together and drove him off, at real cost to both -- a cost the record does not let me minimize, because a Prince held off by two demigods acting in concert is not a threat any single mortal captain should imagine himself equal to alone.")
caveat("A note on timing, because I do not trust a coincidence I cannot explain and I will not hand a cadet a false certainty in place of an honest gap: the Reman line's fall to an Akaviri betrayal is recorded as happening within the same span. Whether the two events are connected, whether the assault on Mournhold was arranged to draw eyes from a capital where a different knife was already being sharpened, I do not know, and no source I have found settles it. I record the proximity and leave the conclusion to better historians than myself.")
body("What I will not leave open to interpretation: the grievance came first, in both directions -- a Duke's cruelty, a petitioner's revenge -- and Dagon arrived only once that grievance had already curdled into something a mortal was willing to sell his own city for. He is not the author of the wound in this case any more than in the last. He is what a wound left open long enough eventually calls in.")

# =====================================================================
kicker("The third case (my own institution's)")
heading1("The Battlespire, and the Lesson I Was Standing Inside Of")
body("I did not read this account. I was there, and I have spent every year since teaching from what it cost us, which is the only honest qualification I hold for writing any of this book at all.")
body("An imperial official, having usurped the seat of an absent Emperor through means this college's records treat as settled fact rather than rumor, made his own bargain with the Prince of Revolution, and the price of that bargain was this academy. We were the Empire's training ground for the mages who might have opposed him, and so we were the first thing he was sent to cripple. He did not need to conquer the Empire outright. He needed only to remove the one institution capable of teaching the Empire to defend itself, and an ambitious, frightened usurper handed him the invitation free of charge.")
body("I will not describe the assault itself at length. Every cadet who sits in this hall studies its tactical record in a separate course, and I am not the right voice for that lesson. I will say only what belongs in this book specifically: the usurper did not summon a besieging force because he wanted this college destroyed for its own sake. He wanted the Empire weak enough that his own seizure of it would hold, and Dagon's Legions were the fastest tool available to make it so. The pattern from Lig and from Mournhold repeats exactly. Someone with power he had not earned needed an injustice enforced quickly, and reached for the Prince who specializes in exactly that transaction.")
doctrine("The lesson this cost us directly",
    "An institution's strength is not measured by its walls. It is measured by how much internal rot exists for an outside force to be invited in through. We did not fall to Dagon's Legions. We fell to the ambition of one of our own Empire's officials, who found Dagon a willing and inexpensive contractor.")

# =====================================================================
kicker("The fourth case (within living memory of this hall)")
heading1("The Oblivion Crisis")
body("The most recent case, and the one every cadet in this room can find a grandparent who remembers directly, so I will be brief where the general record is already thorough and dwell only on what a soldier should extract from it.")
body("A cult calling itself the sons of the new dawn, working from a text they held to be the Prince's own scripture, engineered the assassination of an Emperor and the death of every legitimate heir they could locate, in pursuit of a single goal: the darkening of the old covenant that had kept this Prince and his kind sealed away from Mundus for an age. They very nearly succeeded. Gates opened across the whole of Tamriel at once, and for a season this Empire fought a war on its own soil against an enemy that did not need ships or roads to arrive.")
body("It was ended, in the end, not by armies, which held the line but did not win it, but by a single bastard claimant to a broken throne who gave up his own life to close the door personally. I have made my peace with teaching that fact plainly to cadets who would prefer a more soldierly ending. The war college did not save the Empire that year. One exhausted young man with nothing left to bargain did, and any doctrine that pretends otherwise is vanity, not instruction.")
caveat("Cadets sometimes ask why the cult succeeded as far as it did, given how absurd its scripture reads on a cold morning in this hall. I have no comforting answer. A line of succession thin enough, a court complacent enough, and a population tired enough of both will listen to nearly anything that promises the exhaustion has a cause and an end. That is not a flaw in the Empire that year. That is a flaw in every empire, in every year, and this Prince has never once failed to find it when it was there to find.")

# =====================================================================
kicker("On his instruments, briefly")
heading1("The Tools He Has Left Behind")
body("A war college concerns itself with what can be carried and what it does. I will not repeat what the theological texts already catalogue at length -- his book, his dagger, his gates -- and will note instead only what those texts leave thin.")
body("The Daedric Crescents, curved blades once carried in numbers by warriors sworn to his revolts, were mostly gathered up and destroyed when this Empire reclaimed the Battlespire from the wreckage of its fall; only the rare survivor is now known to exist, and I have not personally handled one.")
caveat("One further weapon is attributed to him in the older accounts I hold, under a name I do not trust my own sources to render correctly -- I have seen it given as something like a sword named for a moon-touched river, though I suspect that is a corruption of an older and more accurate name I have not been able to recover. I record its existence and its reputed power -- strong enough, by the account, to have banished the Prince himself from a battlefield once -- and decline to commit to a spelling I cannot verify. A cadet who finds better sources than mine should correct this entry, not trust it.")

# =====================================================================
kicker("Closing doctrine")
heading1("What a Cadet Is Meant to Take From This")
body("I have taught this course for longer than some of you have been alive, and I will close it the same way every year, because the conclusion has never once needed revising.")
body("You will be tempted, studying this Prince, to think of him as a weather system -- a disaster that arrives, does its damage, and departs, answerable to nothing a mortal does. I have shown you four cases across every age this Empire has kept records for, and in every one of them a mortal hand lit the fire he arrived to burn with. A tyrant's cruelty at Lig. A Duke's cruelty and a petitioner's grudge at Mournhold. A usurper's ambition at this very Battlespire. A court too thin and too complacent during the Crisis. He did not create a single one of those conditions. He has simply never once failed to notice when they were present, and he has never once declined the invitation.")
doctrine("The final word",
    "You cannot build a wall this Prince cannot eventually find a way through, because the wall was never the actual defense. The defense is not giving him a reason to come. Everything else in this book is detail.")
attribution("-- an Instructor, the Imperial Battlespire")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Red Legions.docx")
doc.save(out_path)
print("SAVED:", out_path)
