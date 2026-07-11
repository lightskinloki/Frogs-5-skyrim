# -*- coding: utf-8 -*-
"""Builds 'The Dawn Is Yours' — a recovered Mythic Dawn recruitment tract in the
in-cult voice, annotated in the margins by the Blades agent who confiscated it.
Styled .docx, ember-and-steel palette."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Mythic Dawn: ember red, rust, Gaia Alata gold — vs. Blades steel) ----
EMBER   = RGBColor(0x8E, 0x1B, 0x14)   # ember red (headings)
RUST    = RGBColor(0x6E, 0x2A, 0x16)   # burnt rust (subheads)
GOLD     = RGBColor(0x8A, 0x63, 0x1B)  # Gaia Alata gold (accents in-text)
INK     = RGBColor(0x1E, 0x1A, 0x18)   # warm near-black body ink
STEEL   = RGBColor(0x35, 0x44, 0x50)   # Blades steel (marginalia ink)
HEADBG  = "EEDCD6"   # warm ember-grey table header
MARGBG  = "E9EAEC"   # cool steel-grey (Blades note bg)
BORDER  = "C4A9A0"
ACCENT  = "9A6B1E"   # gold accent for rules/borders
STEELBAR = "35444F"  # marginalia left border
BODYFONT = "Cambria"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

# ---- page = US Letter, 1in margins ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

# ---- helpers ----
def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = EMBER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RUST
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def scripture(text):
    """An indented, gold-italic quotation from the Mysterium Xarxes or the Commentaries."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = RUST
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=RUST, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def epigraph(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11.5); r.font.italic = True; r.font.color.rgb = RUST
    return p

def blades_note(label, text):
    """A cool steel-grey callout with a thick steel left border — the Blades agent's hand."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=STEELBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=STEELBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = STEEL
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = STEEL
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=EMBER)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(110)
r = tp.add_run("THE DAWN IS YOURS")
r.font.name = BODYFONT; r.font.size = Pt(32); r.font.bold = True; r.font.color.rgb = EMBER

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(18)
rs = sub.add_run("A Primer for the Novitiate of the Mythic Dawn,\nDrawn from the Commentaries of Mankar Camoran\nupon the Mysterium Xarxes")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = RUST

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(54)
ra = attrib.add_run("Come slow and bring four keys.")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = EMBER

warn = doc.add_paragraph(); warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
warn.paragraph_format.space_before = Pt(60)
rw = warn.add_run("[ Item recovered from a Mythic Dawn cell. Confiscated, read in full, and\n"
                  "annotated for the archive at Cloud Ruler Temple. The tract is left intact\n"
                  "on purpose — know the enemy in his own words. The notes in the margin\n"
                  "are mine. Do not mistake the one voice for the other. — a Blade ]")
rw.font.name = BODYFONT; rw.font.size = Pt(10.5); rw.font.italic = True; rw.font.color.rgb = STEEL

doc.add_page_break()

# =====================================================================
# I. TO THE ONE WHO READS
# =====================================================================
heading1("I.  Greetings, Novitiate")
body("If this book has come into your hands, it did not come by chance. Every quarter has known us, and none bore our passing except with trembling. Perhaps you came to us through war, or study, or shadow, or the alignment of certain snakes. Each path matters in its kind, and the prize is always the same. That you are here at all means you have the worthiness of kings.")
body("Mankar Camoran was once as you are now — asleep, unwise, protonymic, knowing no name but the one his mother's blood gave him. He is no more. He sits now and waits to feast with you upon all the worlds of this cosmos. This book is your door. Read it slowly, for Lord Dagon would have only those clever enough to pause; all else the Orus claims in their full running. The impatience you feel is your first slave to behead.")
scripture("“When I walk the earth again, the faithful among you shall receive your reward, to be set above all other mortals forever. As for the rest, the weak shall be winnowed, the timid cast down, and the mighty shall tremble at my feet and pray for pardon.”  — the Mysterium Xarxes, written by the hand of Lord Dagon himself")
blades_note("a Blade, in the margin",
    "Mark the shape of it before the theology starts: flattery first. “Worthiness of kings,” “chosen,” “no chance.” Every man who ever put on the red robes was told he was special on the first page. Uriel Septim is dead because men wanted to believe this paragraph. Keep that in front of you as you read the rest, because the rest is better written than this, and that is the danger.")

# =====================================================================
# II. THE LIE
# =====================================================================
heading1("II.  The Lie You Were Born Into")
body("You have been taught that the Divines are your makers and your keepers — Akatosh the Dragon and his eight kin, who raised the world and hold it against the dark. Ask yourself the questions the priests will not answer. How is it that mighty gods die, yet the Daedra stand incorruptible? How is it that the Daedra forthrightly proclaim themselves to men, while your gods cower behind statues and the faithless words of traitor priests?")
body("It is simple. They are not gods at all. The truth has stood in front of you since you were born. The Daedra are the true powers of this universe — the first causes, the fundamental forces of fate, night, and destruction. What are scholarship, love, and mercy beside these? Julianos, Dibella, Stendarr, and all their kin are trifling shadows of first causes. They are Lorkhan's betrayers, posing as divinities in a principality that has lost its guiding light.")
body("Hear the true history the traitor-priests rewrote. Lorkhan was no trickster. He was a Prince of Oblivion, and the world you call Tamriel was his own realm — one more principality among the shining gems of the black reaches, no different in kind from the Coldharbour, from the Quagmire, from the Moonshadow. His own servants turned on him. They tore out his heart, raised the Adamantine Tower to lock the laws of their theft in place, and made themselves gods over the wreckage. They sundered every mortal from the divine spark that was his birthright and yours. That severance is what you call mortality. It is not the order of the world. It is a crime still in progress.")
blades_note("a Blade, in the margin",
    "Two things. First: he lists the Daedric realms to sound learned, and gets three wrong in a row — Coldharbour is Molag Bal's, not Meridia's; the Quagmire is Vaermina's, not Peryite's; Moonshadow is Azura's, not Mephala's. A man who cannot name the Princes correctly is guessing at the rest. Second, and heavier: I put this question to the Dremora Lyranth, who has no love of mortals and no reason to flatter us. Her words: “Mehrunes Dagon's association with Nirn is more metaphorical than existential. Why does Dagon believe Nirn belongs to him? A better question is why do the people of Nirn believe Dagon belongs to them.” The whole tower this book is built on — that Mundus is a stolen Oblivion realm — a Daedra herself calls a mortal's conceit.")

# =====================================================================
# III. THE COVENANT OF CHAINS
# =====================================================================
heading1("III.  The Covenant of Chains")
body("The jailers were not content to steal the world once. They built a lock and gave the key to a bloodline. When the Slave Queen Alessia rose against the Ayleids, she went not to freedom but to a subtler master: she knelt to Akatosh, and he drew a burning handful of his heart's blood and made a covenant of it. So long as her heirs wore the Amulet of Kings and kept the Dragonfires lit in the Temple of the One, the barrier between your world and the freedom beyond it would hold fast, and no great host of the liberated could cross into Mundus.")
body("Understand what the priests call a blessing. The Dragonfires are not a hearth. They are the bars of the cage, kindled anew at the neck of each emperor, kept burning so that the true powers cannot reach the ones they would set free. The red diamond at the throat of every Septim is the lock made jewel. This is why the covenant had to be broken, and why the breaking began with the death of a single old man in a tunnel beneath his own city. With no heir crowned and no fire lit, the barrier thinned to nothing, and the gates could open at last.")
scripture("“Should the Dragonfires fail, and should no heir of our joint blood wear the Amulet of Kings, then shall the Empire descend into darkness, and the Demon Lords of Misrule shall govern the land.”  — the Trials of Saint Alessia. They wrote it as a warning. We read it as a promise.")
blades_note("a Blade, in the margin",
    "Here is what “the breaking began with one old man” actually means, written plainly so no novitiate can hide behind the poetry: they murdered Uriel Septim and all three of his known sons in a single stroke, and when they could not find the last heir, they opened a Great Gate over Kvatch and put a whole city to the flame on the chance the boy was somewhere in it. That is the “liberation” this chapter is selling. Every soul in Kvatch died so this book could call a lock a cage.")

# =====================================================================
# IV. MEHRUNES DAGON, THE LIBERATOR
# =====================================================================
heading1("IV.  Mehrunes Dagon, the Liberator")
body("The world has been taught to fear Lord Dagon as a demon of ruin. That fear is the jailer's last defense, for the one power that can break a prison is the one they must teach you to dread. He is not the Prince of senseless destruction. He is the Prince of ambition, of revolution, of the necessary fire. He does not come to invade Tamriel. He comes to reclaim it — the birthright of a betrayed lord — and to liberate the occupied lands.")
body("This is not the first time he has done such a work, and the Commentaries preserve the memory of it. In the adjacent place called Lig, an old and folded world, mortals lived beneath the Dread Kings, tyrants who owned even the oceans and suffered no soul to live outside their sufferance. The Magna Ge, working in secret in the very bowels of that world, made a Prince of good and poured into him oblivion's rarest asset — hope. That Prince was Mehrunes Dagon. He threw down Lig and cracked its face and declared its every ocean free, and the slaves took chains and teeth to their jailers, and all hope was brushfire.")
scripture("“So shall he crack the Serpent Crown of the Cyrodiils and make federation. There is no dominion save free will. Mehrunes is come. Let all the Orbis know itself to be free.”  — the Commentaries, Volume the Fourth")
blades_note("a Blade, in the margin",
    "“The milk of the unenslaved,” he calls destruction elsewhere. Nourishment. Read the Fourth Volume to its end and it is a list of cities thrown down and peoples crushed — Kyne thrown down, Y'ha thrown down, Horma Gild ground to salt — and every atrocity dressed as a freeing. That is the whole trick of this order: it never once names a cruelty as a cruelty. Learn to hear “liberation” and see the ash under it, or this book will do to you what it did to its author.")

# =====================================================================
# V. NUMANTIA
# =====================================================================
heading1("V.  Numantia, Which Is Liberty")
body("Here is the heart of everything, the one word the whole Dawn turns upon: Numantia. Render it liberty, but do not think it a small liberty — not merely the freedom to move or to choose. It is freedom from history itself. Freedom from the identity that birth pressed upon you before you could refuse it. Freedom from your parents, your ancestors, your culture, the stars that stood over your cradle and dealt you a fate you never chose. You did not choose to be born as you are. Numantia is the promise that you need not remain so.")
body("Consider honestly how much of you was decided without your leave. Your blood, your station, your beginning, the whole determined weight of what came before — the Aedra call this the order of creation, and bind you to the turning wheel with it. We call it a sentence passed on a prisoner who was never tried. In the garden of the Dawn you shall be royalty, a new breed, your own maker. You shall return to your first primal whale and yet come out different — master akin to master, whose mother is miasma; that is, no one, nothing, unbirthed, free.")
scripture("“Palace, hut, or cave, you have left all the fog worlds of conception behind. Numantia, liberty. Rejoice in the promise of paradise — endlessly it shall form and reform around you.”")
blades_note("a Blade, in the margin",
    "I will grant this chapter its due, because pretending it has no pull is how it wins. The ache under Numantia is real. Anyone born into horror has wished they could unwrite it. That is exactly why it is dangerous — a true wound with a poisoned cure poured into it. The Dawn does not offer to heal the wound. It offers to burn down the world that made it, and calls the fire freedom. A man who would murder an emperor to escape his own birth has not become free. He has only found a grander name for his rage.")

# =====================================================================
# VI. THE MYSTERIUM XARXES AND THE FOUR KEYS
# =====================================================================
heading1("VI.  The Book and the Four Keys")
body("All of this was written before you, before all of us, by Lord Dagon's own hand in the Deserts of Rust and Wounds. Its name is the Mysterium Xarxes — Aldmeris Et'Ada Aggregata, forefather to the wife of all enigma. Each word is razor-fed and secret, thinner than cataclysms, tarnished like red drink. It is older and truer than any tome the traitor-priests will show you; the very Oghma Infinium is but its lesser child.")
body("The book does not merely teach. It opens a way. Within it are set the means to build the door to Paradise, and the door has four keys — named in the poet's tongue, that only the worthy might read them:")
scripture("“In his first arm, a storm. His second, the rush of plagued rain. The third, all the tinder of Anu. And the fourth, the very eyes of Padomay.”")
body("These are not riddles for their own sake. Each key is a thing in the world, to be gathered by a hand that has learned to read the Xarxes rightly — a vessel of storm-caught magic, a rain of ruin loosed through an opened gate, the kindling drawn from the light of the Aedra themselves, and an eye of the Padomaic dark torn from a Prince's own hoard. Bring the four, come slow, and enter as Lord Dagon has written.")
lore_table(
    ["The key, as the Xarxes names it", "What the worthy gather"],
    [["A storm (the first arm)", "A vessel of great magical energy, storm-caught and starlit."],
     ["The rush of plagued rain (the second)", "The ruin loosed through an opened gate — quick as rain, and destroying."],
     ["All the tinder of Anu (the third)", "The kindling of the Aedra: the blood of a Divine."],
     ["The very eyes of Padomay (the fourth)", "An eye of the Daedric dark: an artifact of a Prince."]],
    [3.0, 3.5])
blades_note("a Blade, in the margin",
    "This is the part that is not philosophy but operations, and the part that cost us the world's one hope for a season. The four keys open a portal to Camoran's paradise, Gaia Alata, where he kept the Amulet of Kings after his people took it from the Grandmaster's own hands. Note well: they took it from Cloud Ruler's reach. The Dawn's true weapon was never Dagon. It was intelligence — a secret passage only the Blades knew, a bastard son only the Emperor knew, and both known to the enemy. Somewhere a mouth was open that should have been shut. Guard yours.")

# =====================================================================
# VII. THE PATH OF THE DAWN
# =====================================================================
heading1("VII.  The Path of the Dawn")
body("Our order is founded upon the principles of Lord Dagon's mighty Razor, and it rises in four degrees: novitiate, questing knight, chaplain, and master. You stand upon the first. Do not mistake the invitation for the address; you have been welcomed, but you have not yet been shown the harbor. That is by design. The full way is earned, word by word, out of the four volumes of the Commentaries — three that a seeking hand may find, and a fourth that comes only through a sponsor already within.")
body("Read them in order and let the impatience burn off you. Where you feel a shadow-choir gather at the edge of the room, where the candlelight seems to grow eyes and whisper — those are the tower-traitors, the star-orphans the priests call gods, come to frighten you back to sleep. Scorn them. Call out their base natures. Fear only for a second; shaken belief is water for a purpose. Even the Usurper went under the Iliac before he rose to claim his feet.")
scripture("“Roaring, I wandered until I grew hoarse with the gospel. My words found no purchase until I became hidden. Know that humility was Mankar Camoran's original wisdom. Come slow, and bring four keys.”")
blades_note("a Blade, in the margin",
    "Four degrees, and a fourth book only a sponsor can give you — that is not scripture, that is how a cell keeps its members from leaving and its enemies from mapping it. By the time a novitiate holds the fourth volume he has been vouched for, watched, and made to inform on the three volumes' worth of men who came before him. The “shadow-choir” he tells you to scorn is the ordinary voice of a conscience that still works. When your own doubt starts sounding like an enemy, that is the robe going on.")

# =====================================================================
# VIII. KIM AND THE TOWER
# =====================================================================
heading1("VIII.  The Secret of the Tower")
body("What is Paradise for, if not to make you as Lord Dagon is? This is the third secret of Numantia, and the deepest: the way by which mortals become makers, and makers may become mortal again. The wheel of creation, with its eight spokes of false gods, is only a wheel while it lies flat. Turn it on its side and it is a tower, and a tower has an apex, and at the apex one may be as he wills — more, be as he was and yet changed, and change the path for all who walk it after.")
body("This is Kim, the secret syllable of royalty: the ascension the covenant was built to deny you. The bones of the wheel are the eight, but the bones need their flesh, and that flesh is mankind's own heirloom. The dragon-blooded emperors hid this ascension inside six thousand years of ethereal labyrinth and called the world the Arena, that no one else might climb. Tiber Septim climbed it — he did the thing, he said the thing, and reshaped the very jungles of Cyrodiil into the land you know. What one usurper did for himself, Lord Dagon opens to all.")
scripture("“The tower touches all the mantles of heaven, and by its apex one can be as he will. This is the third key of Numantia, and the secret of how mortals become makers, and makers back to mortals.”")
body("But do not mistake the apex for a thing you reach by climbing alone. There is an instrument to the ascension, and Lord Dagon named it in his own Book: the Razor, which cuts not flesh but the name beneath the flesh — the true and secret shape a soul was pressed into at birth. To rise, you must let the blade to your own protonymic, in a place where the world was first made thin, and there be re-cut into the shape you choose to be.")
body("And you will be asked to lay something down. A love you have carried. A loyalty. A name you have answered to all your life and thought was you. Do not mourn it, novitiate, and do not clutch it — that is the old growth, and the gardener prunes the old growth so the new may come. It is only the marble chipped away so the shape beneath can stand forth at last. What feels like losing is only making room. You give up what you were to have space for what you will be.")
scripture("“Each word is razor-fed and secret, thinner than cataclysms, tarnished like red drink. Your name is now cut into its weight.”")
blades_note("a Blade, in the margin",
    "Even granting the metaphysics — and the Dremora would laugh at us for granting them — notice who ever actually reaches the apex in every tale he cites. One man. Tiber. The Usurper. Camoran himself, who wears the Amulet of Kings that slips off any neck without the dragon blood, and calls it proof he remade himself. The tower's promise is offered to thousands and kept by one, and the one is always the man writing the book. Numantia for all; godhood for Mankar. Read the degrees again with that in mind.")

# =====================================================================
# IX. OUR FOUNDER
# =====================================================================
heading1("IX.  On Him Who Opened the Way")
body("Of Mankar Camoran the world tells small and frightened stories — that he is an Altmer four hundred years old and contemporary with the first Septim; that he was born instead a Bosmer, son of the Camoran Usurper and a doomed concubine who cried his name in prophecy before he drew breath. Let the scholars quarrel over which cradle held him. He would tell you the quarrel is the point: a man bound by the accident of his birth has one birth to argue over. A man who has taken Numantia has as many as he requires.")
body("Hear only what matters. He came upon the Mysterium Xarxes and was remade by it — not in figure of speech, but truly. He offered himself to that daybreak, and he set Lord Dagon's razor to his own name, and he laid down the Bosmer he had been as a man lays down a coat he has outgrown. After three nights he could speak fire, and when his voice returned it spoke with another tongue. Whatever blood he was born with, he wears now at his throat the red diamond of emperors, which no hand but a dragon's may fasten. He carved himself dragon, and the universe believed the carving. That is not sorcery, novitiate. That is Numantia made flesh: a man who refused the name he was given and cut himself a truer one.")
attribution("— set down by a Chaplain of the Mythic Dawn, that the novitiate might read and rise")
blades_note("a Blade, in the margin",
    "Here is the passage I would burn, if burning it would unmake the thing it describes. Strip the rapture off and it is a method: the Razor, a place where the veil is thin, and a piece of your own soul given up to the blade. I have stood near one who was carved. He could do things a man should not — and he was less than the man who walked in. The “old growth” they tell you to prune away so gladly is the part that made him a person: the love, the grief, the memory that has weight. What fills the space after is not nothing. It is something. It is simply not him. And the worst of it, the thing that should keep any of us awake: cut rightly, the forgery reads true. It would wear the Amulet. It would sound to a keeper of the old blood like the real thing. I know no test that tells the made soul from the born one. Pray the making stays as hard as he makes it sound.")

# =====================================================================
# X. CLOSING
# =====================================================================
heading1("X.  The Garden Waits")
body("Seek your pocket now and look: there is the first key, glinting with the light of a new dawn. Palace, hut, or cave, you have left the fog worlds of conception behind you. When Lord Dagon walks the earth again the faithful shall be set above all other mortals forever, and the garden shall flood with flowers known and unknown, as it was in the mythic dawn. Endlessly it shall form and reform around you. Come slow, and bring four keys.")
scripture("“Now I sit and wait to feast with thee on all the worlds of this cosmos. Numantia, liberty.”")
blades_note("a Blade, closing the file",
    "And there it ends, the way it began — with a promise that costs you nothing to hear and everything to believe. I have copied it whole into the archive because a threat you refuse to read is a threat you cannot fight, and there will be more of these, and more lost men to hold them. If you are of the Order that guards the heir, read this to its last line and remember the smoke over Kvatch. If you are the lost man himself, and this book found you before we did: the garden is a painting of a paradise held up by a murderer who wears a dead king's jewel. Put it down. Close shut the jaws of Oblivion.")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Dawn Is Yours.docx")
doc.save(out_path)
print("SAVED:", out_path)
