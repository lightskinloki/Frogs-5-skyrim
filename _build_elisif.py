# -*- coding: utf-8 -*-
"""Builds 'What I Would Not Say Aloud' -- the private journal of Jarl Elisif the
Fair. Third Dagon-adjacent lore book, deliberately different genre from the other
two: not cult scripture (The Dawn Is Yours), not military doctrine (The Red
Legions), but a genuinely LIVE personal reckoning by someone who has not yet
decided the answer to "is this worth it" when she begins writing, and who arrives
at a real, dark decision by the end. MAJOR CAMPAIGN CANON, confirmed by the GM
directly this session, not found in any prior file:
  - Elisif genuinely, romantically loves Valerius Caelus -- not merely deceived,
    truly falling for him -- while still genuinely grieving her murdered husband
    Torygg, and suffers real guilt over the overlap.
  - She ultimately, by the end of this journal, DECIDES TO JOIN THE MYTHIC DAWN.
  - She has access to Castle Dour's Imperial military archive and the Blue
    Palace's royal library -- legitimate historical/military records of Dagon's
    past incidents (matching the institutional record The Red Legions draws on),
    NOT the actual Mythic Dawn Commentaries openly (too dangerous for a sitting
    Jarl to hold visibly) -- she has to reach further and more secretly for that
    as her curiosity deepens, tracking her escalating commitment across entries.
Thematic anchor, grounded in play (chapter 9, lines 212-220): Elisif told Davinia
that the party represented "Hope... that the province can be saved without blood
and civil war," and confessed fear that Torygg wasn't strong enough to hold the
kingdom together. He is dead; the war he feared is happening anyway; the same
word -- hope -- is also, per both existing Dagon books, the exact thing the
Magna Ge tried to build into Dagon and got a god of devastating change instead.
This journal is Elisif tracing that same collapse in herself, using the
historical Dagon cases (Lig, Kinlady Estre's failed coup -- a political parallel
to her own position -- Silus Vesuius's family history, the Paradise-is-torture
revelation) as the real evidence she is weighing, and arriving at conversion by
the end. Dramatic irony throughout, never stated outright: she is falling for
the literal embodiment of controlling stasis (Valerius) while reasoning her way
toward the god of its opposite (violent change), and never once sees the
contradiction -- the GM/table does.
Single voice, dated private-journal entries (not a treatise, not scripture).
Styled .docx, mourning silver/deep navy palette with one warm ember accent
(Valerius's intrusion into her cold grief) -- distinct from both other books.
Run: python _build_elisif.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (mourning silver + deep navy, ONE warm ember accent for Valerius) ----
NAVY     = RGBColor(0x22, 0x2B, 0x3A)   # deep mourning navy (headings)
SLATE    = RGBColor(0x3E, 0x4A, 0x5C)   # cool slate (subheads)
SILVER   = RGBColor(0x8C, 0x94, 0xA0)   # moonlight silver (accent/rules)
INK      = RGBColor(0x1C, 0x1E, 0x24)   # near-black body ink
FADED    = RGBColor(0x59, 0x60, 0x6C)   # faded grey (margin asides)
EMBER    = RGBColor(0x8A, 0x42, 0x28)   # ONE warm color, reserved only for Valerius
HEADBG   = "DEE2E8"   # cool silver table header
BOXBG    = "E6E9EE"   # entry-date box bg
BORDER   = "AEB6C0"
ACCENT   = "4A566A"   # slate accent for heading rules
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def entry_date(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.bold = True
    r.font.color.rgb = SLATE
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def ember_line(text):
    """A line about Valerius specifically -- the one warm color in the book."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = EMBER
    return p

def margin(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = FADED
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=FADED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(96)
r = tp.add_run("WHAT I WOULD NOT\nSAY ALOUD")
r.font.name = BODYFONT; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("A private book. Not for the court. Not for the archive.\nFor no one but myself.")
rs.font.name = BODYFONT; rs.font.size = Pt(12); rs.font.italic = True; rs.font.color.rgb = SLATE

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("E.")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = NAVY

doc.add_page_break()

# =====================================================================
heading1("First Entry")
entry_date("Late autumn, alone, past the hour I should be asleep")
body("I am writing this in a hand I do not intend anyone to read, in a book I intend to burn or bury or simply never show, because there is nowhere else in this palace where I am allowed to think a thought before I have already decided what it is safe for me to have thought. Every room in the Blue Palace has an ear in it now. This page does not.")
body("Torygg has been dead the better part of a year. I still reach for his side of a bed he did not sleep in half the nights of our marriage because the Jarl of Haafingar does not get to keep ordinary hours, and I still find the reaching, and I still have to remember, each time, that he is not there to be reached for. Grief does not improve with practice. It only becomes a thing you are better at hiding while it is happening to you.")
body("I told a stranger once -- one of the Wardens, a soldier's daughter with a soldier's bearing -- that the word I wanted for Skyrim was hope. That we could be saved without more blood. Torygg worried, always, that he was not strong enough to hold the province together against Ulfric, against the Thalmor, against his own advisors' worse instincts. I told that soldier my husband was a good king. I believed it. I believe it still.")
body("He is dead, and the province is not saved, and there is more blood than either of us feared, and I do not know what to do with a word like hope once I have watched it fail this completely, in front of me, in my own hall.")

# =====================================================================
heading1("Second Entry")
entry_date("A fortnight later")
body("I must write the other thing, because if I cannot say it to this page I will say it to no one, and a thing that is never said to anyone rots in a person differently than a thing that is at least once spoken.")
ember_line("There is a man. Cyrodiilic, cultured, older than his bearing shows, with a claim to blood so ancient it makes my own look recent. He has been patient with me in a way no one at this court has troubled to be patient, not once, not even in the depths of my own mourning, and I find myself thinking of him at hours I ought to be thinking of the war, or the treasury, or my husband's memory.")
body("I feel it as a kind of theft. As though some part of me that ought to belong entirely to Torygg's grief has gone and given itself to a stranger's company instead, and I cannot make the feeling stop by being ashamed of it, which I have tried, at length, and it has not worked.")
margin("A widow is permitted to remarry. No law and no custom forbids it. I know this. It does not touch the actual weight of the thing, which is not legal at all.")
body("I do not know if I love him. I know that when he is in a room, I stop, for a moment, being only a woman drowning in a war her husband died trying to prevent, and I become instead a person he seems genuinely, unhurriedly glad to look at. I have not felt looked at like that in longer than I can honestly place. I am aware that this is precisely the condition in which a lonely widow makes the worst decision of her life. I am writing that awareness down specifically so that I cannot later claim I did not have it.")

# =====================================================================
heading1("Third Entry")
entry_date("Midwinter")
body("Falk is gone from my counsel -- a long illness, they tell me, and I do not entirely believe the shape of it, though I could not say what I believe instead. Sybille keeps her own counsel these days in a way she did not used to. I am aware, in the part of my mind I do not often visit, that the court around me has been thinning for a year in a pattern I have not troubled to examine as a pattern. I do not examine it now either. I set the awareness down and move past it, because I do not currently have the strength to hold both my grief and my suspicion in the same two hands.")
body("What I will examine, because it costs me less to look at directly, is an older question, and I came to it honestly: if hope was not enough to save my husband, and the war he feared is here regardless of every careful, patient thing he did to prevent it, what would have been enough? I do not ask this as despair. I ask it as a ruler is required to ask every question that might actually matter, however uncomfortable the asking.")
body("Castle Dour keeps the old military record of every incursion this Empire has documented from the Deadlands. I have read a portion of it before, in passing, the way one reads the history of a war one assumes will never come again. I am going to read the rest of it properly now, and I am going to be honest in this book, if nowhere else, about why.")

# =====================================================================
heading1("Fourth Entry")
entry_date("A week into the Castle Dour records")
body("I have been reading about a woman named Estre, a highborn lady in the Summerset court a great many years before my own time, who tried to unseat a queen she believed unfit to hold what she held. She went to Mehrunes Dagon for the means. She was given power, briefly, and then her plot was uncovered and her supporters broken, and whatever the Prince gave her, the record does not trouble to say it saved her.")
body("I read this account three times before I let myself admit why I could not put it down. She was not a peasant with nothing left to lose. She was a woman of standing, watching a throne she believed was managed badly, deciding that the ordinary channels of influence and patience were not going to be enough, and reaching for the one power in the cosmos whose entire nature is the overturning of exactly that kind of failure.")
margin("I am a Jarl watching an Empire managed badly by men very far from here, in a war my husband died trying to prevent, surrounded by a court I am no longer certain I can fully trust. I did not choose this comparison. It simply would not stop making itself.")
body("Her failure is supposed to be the lesson. I notice, and record here because I promised myself honesty in this book if nowhere else, that I did not come away from her account thinking that Dagon had failed her. I came away thinking she had asked for the wrong thing, or asked with the wrong hand, or simply been unlucky in the particular unforgiving arithmetic of getting caught. That is not the response a loyal subject of the Nine is meant to have to a cautionary tale. I am recording that it is the response I had anyway.")

# =====================================================================
heading1("Fifth Entry")
entry_date("Late winter")
body("There is a private museum in Dawnstar I have heard spoken of at court, half in mockery, kept by a man descended from one of the agents of the old Mythic Dawn -- the very cult whose crisis nearly ended the Third Era. He is said to display his ancestor's relics with something close to pride, and the court finds this either tasteless or quaint, depending on who is telling it.")
body("I do not find it quaint. I find that I understand him better than I am comfortable admitting. A man builds a museum to a thing his family was part of because the alternative is admitting the thing was only ever horror, start to finish, and some part of him needs it to have been more than that -- needs the sacrifice, however monstrous its cause, to have meant something, so that his own inherited connection to it is a legacy and not simply a shame he carries for nothing.")
body("I understand the needing. I am beginning to suspect I am building a version of the same museum in this very book, one entry at a time, and calling it honesty instead of a shrine.")

# =====================================================================
heading1("Sixth Entry")
entry_date("Early spring")
ember_line("He asked me, gently, whether I was well. I have not been sleeping and it shows on me now in ways paint cannot fully correct, and he noticed, and he did not press, and he simply sat with me until I was ready to speak or not speak, whichever I chose. I chose not to. He did not seem to mind either way. I do not know how to describe what that costs me to receive from someone who is not my husband, except to say that it costs a great deal, and I keep accepting it anyway.")
body("I have gone further into the Blue Palace's own collection now, past the military record and into the political philosophy the library keeps on the old heresies -- not the banned scripture itself, which no sane Jarl keeps under her own roof, but the commentaries written about it by scholars safely on the correct side of history. Even filtered through disapproval, the actual argument comes through if a reader is willing to hear it plainly instead of only hearing the disapproval wrapped around it.")
body("The argument, stripped of the horror the scholars pile onto it for safety's sake, is this: that a world arranged so that the same families stay comfortable and the same families stay crushed, generation after generation, calling the arrangement peace, is not actually at peace. It is only quiet. And quiet is not the same thing as good, however much easier it is to govern.")
margin("Torygg governed for quiet. I do not say this to condemn him -- I do not believe there was a better available choice, and I loved him for trying regardless. I say it because I am starting to suspect quiet was never going to be enough, and no one at that court, myself included, was willing to say so while he was alive to hear it.")

# =====================================================================
heading1("Seventh Entry")
entry_date("Late spring, and I nearly did not write this one down")
body("I found the darker part of the record today, the part the scholars mention only to condemn and never to explain in full, and I will set it down here plainly because I promised this book my honesty and I am not going to break that promise now that it has become uncomfortable to keep.")
body("The cultists of the old Mythic Dawn were promised a paradise of their own -- a realm built by their prophet from the very text of their scripture, where they would be rewarded eternally for their devotion. What the record actually says happened to them there, once the doors of that paradise closed behind them, is not reward. It is torment, endless and inescapable, dressed in the shape of the garden they were promised. They cannot even die of it. They simply go on.")
body("I sat with that for longer than I am willing to write down the exact count of hours. I expected it to end the question for me. I will record, because I said I would be honest, that it did not end the question. It changed the shape of the question instead. It is no longer, for me, whether Dagon rewards his faithful honestly -- he plainly, demonstrably does not, and I have no illusions left on that count. The question that remains, and I do not like that it remains, is whether the reward was ever truly the point for me in the first place, or whether what I actually wanted was simply for someone, something, some force in this cosmos, to finally act instead of merely enduring.")

# =====================================================================
heading1("Eighth Entry")
entry_date("Summer -- the last entry I intend to make in this particular book")
ember_line("He told me, last night, that he loves me. I believe he meant it. I have decided that I love him also, and I have stopped, as of tonight, apologizing to this page for it.")
body("I have also decided the other thing, and I find, now that it is decided, that I do not need to write pages of justification for it the way I have written for every smaller step that led here. That is, I think, what a true decision feels like, as opposed to a doubt being carried a while longer. Torygg died trying to hold a quiet in place that was never going to hold regardless of what either of us did. I do not intend to spend what remains of my own life holding the same failing quiet in his memory, out of loyalty to a peace that was already gone before he was.")
body("I do not know yet what it will cost me to actually act on this, or how, or when, or in what order I will need to do the things I now believe I am going to do. I know that I came to this book with a question, honestly asked, and that I have answered it, and that the answer is not the one the court, the temple, or my husband's memory would have wished for me. I am not certain it is the answer I would have wished for myself, a year ago, before any of this happened to me. It is the one I have. I intend to act on it.")
body("This book ends here. What follows, follows elsewhere, and not on any page I intend anyone, ever, to find.")
attribution("-- E.")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "What I Would Not Say Aloud.docx")
doc.save(out_path)
print("SAVED:", out_path)
