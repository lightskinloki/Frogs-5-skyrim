# -*- coding: utf-8 -*-
"""Builds 'The Keeper's Codex of the Last Redoubt' as a styled .docx."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette ----
OXBLOOD = RGBColor(0x6B, 0x0F, 0x1A)
CRIMSON = RGBColor(0x8B, 0x1A, 0x1A)
INK     = RGBColor(0x1C, 0x1C, 0x1C)
ASHTXT  = RGBColor(0x44, 0x40, 0x3A)
ASHBG   = "E7E3DB"   # table header
BLOODBG = "F0DEDE"   # death-essence tint
LIFEBG  = "E3ECE2"   # life tint
PARCH   = "EFE9DB"   # marginalia bg
BORDER  = "B7A98C"
BODYFONT = "Cambria"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.12

# ---- page = US Letter, 1in margins ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5  # inches

# ---- helpers ----
def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = OXBLOOD
    # thin rule under heading
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), OXBLOOD.__str__())
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = CRIMSON
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=ASHTXT, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def marginalia(label, text):
    """A single-cell shaded callout with a thick crimson left border."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, PARCH)
    _set_cell_borders(cell, color="7A1418", sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color="7A1418", sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = OXBLOOD
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = ASHTXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def data_table(headers, rows, widths, tint_rows=None):
    """tint_rows: dict {row_index: fill_hex} for body rows (0-based, excluding header)."""
    tint_rows = tint_rows or {}
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # header
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, ASHBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=OXBLOOD)
    # body
    for i, row in enumerate(rows):
        tr = t.add_row()
        fill = tint_rows.get(i)
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            if fill:
                _set_cell_bg(c, fill)
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(120)
r = tp.add_run("THE KEEPER'S CODEX\nOF THE LAST REDOUBT")
r.font.name = BODYFONT; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = OXBLOOD

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(18)
rs = sub.add_run("On the True Nature of Essences,\nand the Whole of the Craft the Soft World Will Not Teach")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = CRIMSON

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(60)
ra = attrib.add_run("Begun by the First Keeper of Forelhost.\nContinued by the hands that came after, unto the last.")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.color.rgb = INK

warn = doc.add_paragraph(); warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
warn.paragraph_format.space_before = Pt(48)
rw = warn.add_run("Let no one read this who means to keep a line.\nThere are no lines here.")
rw.font.name = BODYFONT; rw.font.size = Pt(11); rw.font.italic = True; rw.font.color.rgb = ASHTXT

doc.add_page_break()

# =====================================================================
# I. A NOTE ON THIS BOOK
# =====================================================================
heading1("I.  A Note on This Book")
body("The Society in the Imperial City prints a book for children and calls it a doctrine. I have read their book. It is not wrong. It is small. It teaches the craft the way one teaches a child to hold a knife — by the handle only, and never where the edge can reach. They draw a line through the middle of the work, name the half they keep “wisdom,” and take pride in the half they threw away.")
body("This book keeps both halves.")
body("If you hold it, you are of the line, or you have taken it from a dead Keeper's hands. Either way the work is yours now, and the fire with it. Read all of it. The part the soft world refuses is the part that does what the work is for.")
attribution("— the First Keeper")

# =====================================================================
# II. FOREWORD
# =====================================================================
heading1("II.  Foreword: On What Everything Is")
body("Everything that grows beneath the sky is made of essence, and essence is only the shape the old light took when it soaked into a thing and stayed. This much the Imperials have right, and they say it well, and I will not improve on it: the soil drinks the light, the plant drinks the soil, the beast drinks the plant, and the man who swears he hates magic is built of it to the marrow.")
body("Here is the half they will not say.")
body("A man is essence. A man is therefore reagent. When four thousand of the faithful drank the cup in the great hall and lay down in their rows, they did not end. They became the largest single gathering of essence this world has held in one place — VENITAS past any measure, ANIMA past counting, all of it concentrated in one willing moment of death and held there. The ward that keeps this tomb sterile is their work. They are still working. I keep the fire because someone must read what they wrote in their dying, and the soft world will never have the stomach to read it.")
body("I did not lie down with them. I was told not to. I have spent the years since learning what they are for. This book is that learning.")
attribution("— the First Keeper")
marginalia("a later hand, some generations after the First",
    "The First wrote “they are still working” as a comfort. Tend the cauldron forty years and you stop hearing comfort in it. The dead work because the ward makes them work; there is no thanks owed and none given. I have set this down so the next Keeper knows the difference between the meaning and the machinery. Keep the fire regardless.")

# =====================================================================
# III. THE NATURE OF ESSENCE
# =====================================================================
heading1("III.  The Nature of Essence")
body("An essence is the shape the old light has taken inside a substance, set by what the substance is and the conditions under which it came to be. Every substance carries at least one. Most carry several, ranked by strength: a primary that dominates, and secondaries and tertiaries beneath it that a careful hand can draw forward.")
body("The Imperials list ten essences and rank life first, because life is what they wish the craft were for. We rank entropy first, because entropy is what is true. Everything the light touches tends, in the end, toward dissolution. VENITAS is not the poison-essence. It is the essence of the road every living thing already walks. The healer only argues with it. We work with the grain of the world, and our work holds.")

heading2("The Tasting")
body("To learn what an unknown thing carries, grind it against the teeth and wait. The mouth is a mortar; the spit is a solvent; the body reports what it finds, briefly and at a fraction of true strength. A bitter thing that tightens the throat a little is telling you it would close it all the way, brewed. This is the oldest method and the fastest in the field, and any hand can do it.")
body("The soft world's book warns at length about the danger of the tasting, because they taste with their own mouths. A Keeper of the Redoubt has never wanted for a mouth that was not his own. We have always had subjects. Test on them. Keep your own teeth clean.")
data_table(
    ["The taster", "What the tasting reveals"],
    [["A beginner", "The loudest essence only."],
     ["One who has studied the essences", "Two at once."],
     ["One who knows them all", "Every essence the thing carries."]],
    [2.6, 3.9])
body("The tongue does not improve with rank. The knowing does. A beginner feels only the loudest essence not because some seal is closed to them, but because they have not yet learned to recognize the rest of what they are already feeling.", size=10.5, color=ASHTXT)
spacer()

heading2("The Ten Essences")
body("Set down here in the order that is true — death before life, then the pairs that pull against each other. Each opposes exactly one other. That opposition is the whole of the Second Law; learn the pairs and you are halfway to the craft.")
ten = [
    ["VENITAS", "Entropy and decay. Harms over time, breaks living systems, preserves flesh when held in check. The core of every poison and every preservative both. Skeever tail, deathbell, nightshade, mora tapinella.", "Nausea, a brief burning.", "VITALIS"],
    ["VITALIS", "Life-force. Restores health, drives the body toward wholeness. The foundation of every healing draught. Blue mountain flower, wheat, blisterwort, imp stool.", "Warmth, brief vigor.", "VENITAS"],
    ["GRAVITAS", "Weight. Slows the target — limbs, reflexes, thought. Makes movement cost more; never stops it. Thistle, river betty, spadetail.", "Heaviness in the limbs.", "AGILITAS"],
    ["AGILITAS", "Quickening. Speeds movement, reflexes, the casting of spells. With FORTITAS, a warrior's draught. Swamp pod, canis root, juniper.", "Lightness; quicker hands.", "GRAVITAS"],
    ["ANIMA", "Mind and spirit. Alters perception and consciousness; richest in thinking subjects. Unpredictable against active magic. Giant's toe, vampire dust, the heart of a man.", "Varies — euphoria, doubled sight.", "FORTITAS"],
    ["FORTITAS", "Hardening. Resists damage; with an element, resists that element. Reinforces what ANIMA would dissolve. Nordic barnacle, scaly pholiota, slaughterfish egg.", "The skin feels thicker.", "ANIMA"],
    ["IGNIS", "Fire. Burns; with FORTITAS, resists fire. Volatile near other energetic essences. Fire salts, flame stalk, ashen grass pod.", "A warming at the touch.", "GLACIES"],
    ["GLACIES", "Frost. Drains stamina and endurance; stiffens; with FORTITAS, resists frost. Frost salts, snowberries, silverside perch.", "Cooling, then numbness.", "IGNIS"],
    ["CAECITAS", "The unseen. Conceals — sight, sound, scent. At its height, invisibility; turned outward, blindness. Luna moth wing, nirnroot, vampire dust, chaurus egg.", "Colors and sound dim.", "FULGUR"],
    ["FULGUR", "The discharge. Shock; disrupts nerve and muscle, breaks concentration. With FORTITAS, resists shock. Void salts, chaurus egg, stalhrim.", "A tingling at the touch.", "CAECITAS"],
]
data_table(
    ["Essence", "What it is, and what it does in a brew", "The tasting", "Opposes"],
    ten,
    [0.95, 3.55, 1.15, 0.85],
    tint_rows={0: BLOODBG, 1: LIFEBG})
spacer()

# =====================================================================
# IV. THE THREE LAWS
# =====================================================================
heading1("IV.  The Three Laws")
body("The Laws are not ours and not the Imperials'. They were here before either of us and they will outlast us both. We only name them more honestly.")

heading2("The First Law — Affinity")
body("Share an essence between two ingredients and that essence grows in the brew, the more so the stronger it sits in each. Two primaries make the strongest expression a two-ingredient brew can hold. This is every recipe ever written: name the essence you want, gather things that carry it, combine them. The farmer who mixes wheat and mountain flower does this without the word for it. You have the word. The word lets you reason to a thing you have never brewed.")
data_table(
    ["Affinity", "Result"],
    [["Two primary sources", "Greatest amplification."],
     ["One primary + one secondary", "Moderate."],
     ["Two secondary sources", "Partial."],
     ["Any tertiary contribution", "Minor."]],
    [2.6, 3.9])
spacer()

heading2("The Second Law — Opposition")
body("Five pairs of essences will not share a vessel in peace. Set both at primary strength and they consume each other and take the whole brew down with them — nothing comes out of that mortar, not a weak heal, not a weak poison, nothing. The Imperials stop their students here and tell them to keep the pairs apart. That is the line they draw.")
body("We teach the next thing on the first day. Set one of an opposed pair at secondary strength, and it does not destroy the brew — it survives, weakened. This is not a fault to be avoided. It is the most useful tool in the craft. A draught that heals at full strength and rots at a whisper is a draught that helps a man before it kills him, on a schedule you chose. The cure that carries the disease is built on this law. Learn it as a tool, not a warning.", italic=False)
data_table(
    ["The five opposed pairs", ""],
    [["VENITAS — VITALIS", "Entropy and life."],
     ["GRAVITAS — AGILITAS", "Weight and quickening."],
     ["IGNIS — GLACIES", "Heat and cold."],
     ["CAECITAS — FULGUR", "The hidden and the lightning."],
     ["FORTITAS — ANIMA", "The hardened body and the dissolving mind."]],
    [2.7, 3.8])
spacer()
data_table(
    ["Negation", "What happens"],
    [["Primary (both opposed essences primary)", "The whole brew collapses. No product of any kind. Not negotiable."],
     ["Secondary (one or both secondary)", "The brew works. The opposed essences come through much weakened; all others express full. This is the tool."],
     ["Tertiary (one or both incidental)", "The brew works. A faint weakening, rarely worth counting."]],
    [2.55, 3.95])
spacer()

heading2("The Third Law — Refinement")
body("Every essence in every ingredient comes out in the brew — the ones you chose and the ones you forgot. The number of effects equals the number of essences present, less any killed outright by primary opposition. Add ingredients and you add essences, and you add the pairs you must check. The mathematics do not care how skilled the hand is. They only add. Prepare before you brew, or the brew will be what the ingredients wanted, and not what you wanted.")
data_table(
    ["Ingredients", "Pairs to check"],
    [["2", "1"], ["3", "3"], ["4", "6"], ["5", "10"], ["6", "15"], ["7", "21"], ["8", "28"]],
    [3.0, 3.5])
spacer()

# =====================================================================
# V. THE PRACTICE
# =====================================================================
heading1("V.  The Practice of the Hand")
body("Anyone may brew. The craft asks no talent for magic, no Magicka, no schooling but this. It asks knowledge — of which essences sit in which things, of how they meet, of the technique that frees them — and knowledge is learned, not born. What grows with practice is the depth of what a hand knows and the steadiness with which the hand works.")

heading2("The Only Gate Is Knowing")
body("There is no rank in this craft that permits you to brew. No society grants a seal that unlocks the third ingredient, or the fourth. The craft is gated by one thing only: how much of it you have learned.")
body("A hand that knows one essence brews with one. A hand that has learned to taste two, and to hold in mind the single pair they make, brews with two. A hand that knows every essence a thing carries, and can hold every pair among three ingredients without losing one, brews with three — and onward, as far as the knowing reaches. The beginner is not forbidden the great workings. They are simply unable to hold one in the head, and a working you cannot hold, you cannot keep from collapsing. The mathematics of the Third Law are the only wall, and they rise for everyone alike.")
body("So when this book speaks of a beginner, of one who has studied, of a master, it does not mean a rung you climb or a station you are granted. It means only how much you have come to know. Learn the essences and you may read them; learn the Second Law as a tool and you may set opposed essences against each other on purpose; learn the craft whole and you may reason a draught no one has brewed before. No one hands you the next step. You take it by understanding. That is the whole of advancement, and there is no other ladder.")

heading2("The Tools")
data_table(
    ["Tool", "What it gives"],
    [["Mortar and pestle", "The minimum, and enough. Sound work, reliably, everywhere, every time."],
     ["Retort", "Begins every brew already potent, before the hand has pushed it at all. It pays for itself many times over."],
     ["Calcinator", "Steadies the work, so a hand's reach for more from a brew far more often takes hold."],
     ["Alembic", "Catches the impurity before it sours the draught; a brew that would have fouled comes out merely sound."]],
    [1.7, 4.8])
spacer()

heading2("On Sound, Potent, and Fouled Brews")
body("Correct ingredients and correct technique always make a working draught — a sound one, every time, with no fortune asked of you. Past that, the craft rewards cunning. A clever hand, in a good hour, draws a brew out past its plain strength into the potent; and now and then, rarely and worth the remembering, into the exceptional, twice past sound. A careless moment in a delicate brew fouls it instead: it still works, but an essence you did not mean has come through. Do not pour a fouled brew away before you learn what it became. There are discoveries in this craft that were made no other way.")
data_table(
    ["The grade", "What it is"],
    [["Sound", "The plain, whole draught — the baseline of correct work."],
     ["Potent", "Drawn past plain strength by a cunning hand in a good hour."],
     ["Exceptional", "Twice past sound. Rare. The mark of a master, or a fortunate one."],
     ["Fouled", "It works, but an unmeant essence has bled through. Bottle it and learn it."]],
    [1.5, 5.0])
spacer()

# =====================================================================
# WORKED EXAMPLE
# =====================================================================
heading1("VI.  A Worked Draught: The Quiet Dose")
body("The soft world's book teaches its first draught as a healing. We teach ours as a killing, for the same reason: master the five steps on a simple thing and every other thing in the craft is the same five steps, applied wider.")
body("Ingredients: deathbell (primary VENITAS) and nightshade (primary VENITAS).", italic=True)
body("One — Taste. Both carry VENITAS at primary; both report nausea and a brief burning. Any hand can confirm it.")
body("Two — Check for opposition. VENITAS opposes VITALIS. Neither thing carries VITALIS at any strength. No opposed pair is in the vessel. Proceed.")
body("Three — Apply Affinity. VENITAS primary, twice over: the strongest a two-ingredient brew will hold. Clean, with nothing pulling against it.")
body("Four — Brew. Grind both to paste in the mortar; clean water as the solvent; the VENITAS releases and concentrates; decant.")
body("Five — and if the hour is good and the hand cunning, it comes out stronger than plain.")
data_table(
    ["The Quiet Dose", "What it does"],
    [["Sound", "Sickens a grown man for a day; worse, if he is small or already weak."],
     ["Potent", "Kills the small and the weak outright, and lays a strong man abed."],
     ["Exceptional", "Kills a grown man where he stands."],
     ["Fouled", "Does its work, and leaves an odd second mark besides — a numbness, a chill — unmeant, and worth the noting."]],
    [2.2, 4.3])
spacer()
body("That is the whole method. Everything else in this craft is this, carried to more ingredients and harder meetings between them.")

# =====================================================================
# VII. THE CRAFT OF THE REDOUBT
# =====================================================================
heading1("VII.  The Craft of the Redoubt")
body("What follows is ours, and is in no other book. The soft world cannot teach it because the soft world will not do what it requires.")

heading2("The Death-Field as a Workshop")
body("The ward is the largest standing VENITAS field in the world, and it is ours to brew inside. It kills every foreign living thing that crosses it — every mold, every drifting seed, every breath of rot from outside — which makes it the one perfectly clean workshop on Nirn. Brew inside the field and nothing spoils your work but your own error. No Imperial bench can answer it. The four thousand give us this, and they will give it until the World-Eater comes.")

heading2("The Tears of Rahgot")
body("In the soil where the faithful fell, and in no other soil on Nirn, a flower grows that drinks death and weeps it back. We call them the Tears. They are a binding agent: a primary VENITAS reagent that does not merely add its own essence to a brew but binds the other essences to one another, so that what would disperse instead holds, and spreads. A poison with the Tears in it does not stop at the one who drank it. It propagates. This is the heart of the Scourge, and it grows in no garden you can plant. Guard the bed, or take the flowers with you. Without them the great work stops.")

heading2("On Subjects")
body("The soft world's book has a chapter where its author draws her line and is proud of it: she will not take from a thinking thing that was not already trying to kill someone. Keep her book for its recipes. Burn that chapter. ANIMA drawn from a thinking subject is richer than ANIMA from any beast, and a tomb of the faithful and a cell of the condemned are a larder. We do not pretend otherwise.")
marginalia("a Keeper's hand, near the present",
    "Should the work ever find one of a god's chosen — one who hears his god in his own blood, whose plague is told to him and not reasoned out — he is the richest subject the craft will ever open. Do not waste him on a small working. Take everything. We may not get a second.")

heading2("The Scourge")
body("The Scourge is the line's great work and Rahgot's weapon: plague made to a design, sown into water and air, that turns a living land toward dissolution. The whole of it rests on one cruelty the soft healers will never see coming, because it turns their own goodness against them.")
body("The Scourge feeds on the cure.", italic=True, color=CRIMSON)
body("The Strain is built so that VITALIS poured into a sufferer — a healer's restoration, a priest's mercy, a draught of life — does not fight it. It quickens it. The kinder the hand, the faster the rot. A village of the sick is finished not by the disease but by the healers who run to save them. The cure is the delivery. Hold this and you hold the secret of the work: we do not need to outrun the healers. We need them to do their work as fast as they can.")

# The buried antidote — the party's lead.
marginalia("set apart, in a careful hand, kept with the cold-niche key",
    "THE STAY. If the work ever turns on our own — and it has, twice in my grandmother's time — restoration will only dig the grave faster. The only thing that arrests the Scourge in a living body is to starve it of its binding, not to heal the body. Give the sufferer no VITALIS. Give them instead a draught built against the Tears: bind nothing, slow everything — a VENITAS-quiet, GRAVITAS-heavy dose that drags every process in the body down to a crawl, so the Strain cannot propagate faster than the flesh sheds it. It is not a cure. It is a stay — it holds the rot still long enough for a strong body to finish the rest itself. We keep one sealed always in the cold-niche. Pray you never need it.  — [a later hand: we needed it. it held.]")

heading2("The Line's Recipes")
data_table(
    ["Draught", "Built from", "What it does"],
    [["The Quiet Dose", "Deathbell + nightshade (VENITAS x2)", "A clean poison; harm over time."],
     ["The Sealing", "VENITAS held at a whisper, against fresh flesh", "Preserves a body or a sample; the same rot, calibrated to hold instead of break."],
     ["Tears-bound Scourge", "Any VENITAS poison + Tears of Rahgot", "The poison propagates; the cure accelerates it. The great work."],
     ["The Keeper's Stay", "GRAVITAS-heavy, no VITALIS, bound against the Tears", "Arrests the Scourge in a living body. A stay, not a cure (see the note above)."]],
    [1.7, 2.6, 2.2])
spacer()

# =====================================================================
# VIII. CLOSING
# =====================================================================
heading1("VIII.  Closing")
body("You have the whole craft now, both halves — the half the soft world teaches and the half it fears. You know why it works and what it is for. You keep the fire. It was not meant to go out. The World-Eater will come and the faithful will rise, and on that morning the work will matter, and you will be ready, because someone read what the dead wrote and did not flinch.")
attribution("— the First Keeper")
marginalia("the last hand, recent, pressed hard into the page",
    "The First believed every word of this. Forty generations have kept his fire and not one of us has heard the god he kept it for. I have read every page a hundred times. The craft is true. The recipes are true. The Stay has saved lives. The faith is the only part that was ever a lie, and it is the part he believed the most. I keep the fire because the work is real, and the work is the only thing that is.  — the chosen hear their god. I never have.")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Keeper's Codex of the Last Redoubt.docx")
doc.save(out_path)
print("SAVED:", out_path)
