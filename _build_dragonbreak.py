# -*- coding: utf-8 -*-
"""Builds 'What I Understand of Broken Time' -- Martin Septim's private research
notes on Dragon Breaks, written in the days after the Fall of Kvatch, during the
Oblivion Crisis. NOT Vivec's own voice (his real writing, the 36 Lessons, is
deliberately never reproduced or imitated here -- referenced and paraphrased only,
by design, after the GM correctly vetoed attempting to match Michael Kirkbride's
actual prose). Martin encounters Vivec's teaching on broken time for the first
time AFTER Kvatch, as part of trying to understand what he just lived through --
not prior background reading. His voice, calibrated against real, verbatim
Oblivion dialogue (sourced via web search this session, not guessed):
"I'm just a priest. You're looking for a champion." / "Have you brought help?
We've been trapped here since the daedra overran the city." / "They are stable
portals between our world and Oblivion. They violate the most basic principles
of daedric magic." Short declarative sentences. States a limitation or a fact
once, plainly, and moves on -- NO circling back to comment on the act of writing,
no "not because X but because Y" hedging (an earlier draft this session made
exactly this mistake and was correctly rejected as "Claude's voice, not Martin's").

Content: the Dragon Break concept accurately reconstructed from the GM-supplied
transcript in original prose (Akatosh/time, contradictory-but-equally-true
timelines, "untime," even Moth Priests and the Elder Scrolls failing to parse
it), then the real historical examples -- the Time-Wound at the Throat of the
World (Felldir the Old banishing Alduin, already established campaign canon per
PLANNING SCRATCH), the Battle of Red Mountain (Numidium's first activation, the
contradictory accounts), the Middle Dawn (a thousand-year break; the responsible
cult's exact name is uncertain/STT-garbled in the source and is NOT invented or
guessed at here), Tiber Septim's own use of Numidium, and the Warp in the West.
Closes on Martin's own present, unresolved dread -- Mehrunes Dagon is actively
tearing open Oblivion gates right now, and Akatosh's whole domain is the thing
being violated -- without him predicting his own future godhood, which he has
no way to know yet.
Single voice, styled .docx, ash-and-embercoal palette (a burned city, a
frightened priest) -- distinct from every other book built this session.
Run: python _build_dragonbreak.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (a burned city, a frightened priest: ash-grey + embercoal, no gold) ----
ASHBLK   = RGBColor(0x28, 0x26, 0x24)   # ash-black (headings)
EMBERCOAL= RGBColor(0x5A, 0x36, 0x2C)   # dull embercoal (subheads, sparing warmth)
INK      = RGBColor(0x1D, 0x1C, 0x1A)   # near-black body ink
FADED    = RGBColor(0x5E, 0x58, 0x52)   # faded ash-grey (margin/date notes)
HEADBG   = "E6E1DA"   # ash-parchment table header
BOXBG    = "ECE7E0"   # entry-marker bg
BORDER   = "B6AC9E"
ACCENT   = "4A3E36"   # muted ember-ash accent for rules
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = ASHBLK
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def entry_date(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.bold = True
    r.font.color.rgb = EMBERCOAL
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def margin(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = FADED
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=FADED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(96)
r = tp.add_run("WHAT I UNDERSTAND\nOF BROKEN TIME")
r.font.name = BODYFONT; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = ASHBLK

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("Notes kept in the days after Kvatch")
rs.font.name = BODYFONT; rs.font.size = Pt(12); rs.font.italic = True; rs.font.color.rgb = EMBERCOAL

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("Martin")
ra.font.name = BODYFONT; ra.font.size = Pt(13); ra.font.italic = True; ra.font.color.rgb = ASHBLK

doc.add_page_break()

# =====================================================================
heading1("First Night")
entry_date("The road out of Kvatch")
body("Kvatch burned in a day. I watched it from inside my own chapel, and there was nothing in my prayers that stopped it. I am alive because other people died holding a door I was standing behind, and because a stranger walked into that chapel through the worst of it and did not leave until the rest of us could.")
body("I never got a name. I asked. They did not seem the sort who gives one out easily, or perhaps there simply wasn't time. It has not stopped them staying close since, on this road, without being asked to. I am grateful for it more than I have said aloud.")
body("I do not know why I am writing this instead of sleeping. I think it is because I do not want to close my eyes yet and see it again, and a page is at least something to look at instead.")
body("Jauffre gave me a book two nights ago, from whatever the Blades still keep at Weynon Priory. He did not tell me why. I think he wanted me occupied with something other than my own face in a window. It is a hard book. I am going to try to write down what is in it in my own words, because I do not think I understood it the first time through, and writing forces a man to admit what he actually followed and what he only nodded at.")

# =====================================================================
heading1("What the Book Is")
entry_date("Second night")
body("It is Vivec's own teaching. I will not try to write the way he writes. I could not if I tried, and I would only make a fool of myself and a mess of what he actually meant. I am going to say plainly, in my own words, what I think he is describing, and where I am unsure, I am going to say that too.")
body("He is writing about time breaking. Not slowing. Not stopping the way a heart stops. Breaking, the way a bone breaks, so that it no longer holds one shape but several at once, in the same place, and none of them false.")
margin("I have read that last sentence back four times and I still do not fully believe it. I am writing it down anyway because I think it is what he means, and disbelieving a true thing does not make it less true.")

# =====================================================================
heading1("The Plain Version")
entry_date("Same night, later")
body("Here is the plainest way I can put it. Akatosh is the god who keeps one thing happening after another instead of everything happening at once. That is his gift and his work, and it is so constant that no one thinks of it as a gift at all, the same way no one thanks the ground for staying underfoot.")
body("Sometimes, for reasons Vivec does not pretend to fully explain, that stops holding. When it stops, more than one true thing can happen in the same place, to the same people, and neither one erases the other. A war is fought and also not fought. A man dies and also does not. Both accounts are real. Neither is a lie. They simply happened on time that was no longer agreeing with itself.")
body("Then, eventually, it resolves. Time goes back to holding one shape. But everything that happened during the break stays true, all of it, even where it contradicts itself, because it did happen, on time that allowed more than one thing to be happening. People afterward remember it as a kind of grief they cannot point to a cause for. I understand that part better than I would like to.")
margin("Vivec calls it, as near as I can render it in a plainer tongue, an untime. I think that is the most honest word in the whole book. It is not a bad time or a strange time. It is a period that time itself was not fully present for.")

# =====================================================================
heading1("Even the Moth Priests")
entry_date("Third night")
body("One thing in the book steadied me, oddly, because it meant my own confusion was not a personal failing. Vivec writes that even the Moth Priests, reading the Elder Scrolls themselves, cannot see clearly into a period broken this way. The Scrolls hold every true thing that happened. During an untime, they hold several true things that contradict each other, and a man reading them comes away as confused as I am tonight, no matter how holy his order.")
body("If the Scrolls themselves cannot sort it, I do not feel quite so foolish for needing three nights and a burned city behind me to understand it at all.")

# =====================================================================
heading1("The Wound on the Mountain")
entry_date("Fourth night")
body("Vivec names an old one, small as these things go, and it happened here, in Skyrim, on the mountain the Nords call the Throat of the World. During the war against the dragons, in an age before anything I know how to date properly, a man named Felldir stood against Alduin himself with two other Tongues, and they could not kill him. Felldir used the Elder Scroll he carried and the Thu'um together, and instead of killing the dragon, he threw him forward through broken time, out past his own age and into one much later.")
body("I do not know if the men who did this understood what they were doing to time itself, or only that it worked. Vivec does not say. What he does say is that this is the clearest, smallest example he can point to, and that whoever is meant to face Alduin when he returns will be facing a debt that was moved forward through an untime rather than paid at the time it was owed.")
margin("I find I keep returning to this one more than the others. I could not say why, except that a debt moved forward instead of paid feels uncomfortably close to something I understand about myself right now, though I could not say what.")

# =====================================================================
heading1("The Red Year at Red Mountain")
entry_date("Fifth night")
body("The next one Vivec discusses happened at Red Mountain, during the old war between the Dwemer and the Chimer, in the First Era. The Dwemer had built a thing called the Brass God, meant to make its makers something closer to divine. What happened when they woke it, by every account Vivec has gathered, is that time broke around the battle itself.")
body("This is the part that unsettled me most to read, because the disagreement is not small. Some accounts have the Dwemer King falling to a hero's blade. Others have him falling to an arrow. Others say he did not fall in that battle at all, but later, of wounds, or by another hand entirely. Vivec does not treat these as rival rumors to be sorted into one true account. He treats them as equally true, because the battle did not finish happening on time that agreed with itself. Whichever version a given text tells you, it is not lying. It is only telling you which piece of the untime it was standing in.")
body("The Dwemer themselves were never seen again after that war, by anyone, anywhere. I do not think that is a coincidence, though Vivec does not say so directly, and I will not claim more certainty than he gives me.")

# =====================================================================
heading1("A Thousand Years, More or Less")
entry_date("Sixth night")
body("The largest one Vivec describes I can only half-follow, and I will say so honestly rather than pretend otherwise. A sect within the old Elven religious order forced a break by working a rite atop the White-Gold Tower itself, using an instrument of the Aedra's own making, with the stated aim of prying Akatosh apart from an older, more primal aspect of himself.")
body("What followed, by Vivec's account, is remembered as roughly a thousand years, though he is honest that the number itself may be a convenience the historians settled on rather than a true count, since a period where time was not agreeing with itself is a strange thing to measure with a calendar built to assume that it does. Within that span, by his account, entire regions were conquered and also never conquered. The sky itself is remembered as having changed color depending on whose account you trust. It ended, by means Vivec does not claim to fully understand himself, and time resumed holding one shape, with every contradiction of that thousand years left standing, true and unreconciled, underneath it.")
margin("I confess I read this section twice and understood it less the second time than the first. I am recording that honestly rather than pretending a false clarity I do not have.")

# =====================================================================
heading1("Tiber Septim's Hand")
entry_date("Seventh night")
body("The last of the old examples touches my own family's name, which is not a comfortable thing to write down, but I said I would be honest in this book and I intend to keep that.")
body("Tiber Septim, in the course of building the Empire I am told I am somehow heir to, made use of the same Brass God the Dwemer had built at Red Mountain. Vivec's account holds that when it was used again, in his own conquests, the same fracture followed it did before — contested outcomes, battles remembered as lasting an hour by some accounts and years by others, whole stretches of land remembered as poisoned by the working that are, by other accounts, remembered as untouched. Whatever the truth underneath it, the Empire that resulted was real enough. I am living in the shape of what that Empire became, right now, on this road, running from a city it could not save.")

# =====================================================================
heading1("The Nearest One")
entry_date("Eighth night")
body("The last example in the book is close enough in time that men still living could have witnessed the tail end of it, in the far west, in Hammerfell and the lands around it. The same Brass God again — it seems to be the one thing reliably present whenever this happens, and I do not think that is coincidence, though I am not scholar enough to say why with any confidence. Vivec records that a great many small, warring kingdoms in that region were, at the end of their own untime, simply consolidated into a handful of larger ones, with the losing claims and the winning claims both remembered as having genuinely happened, by different people, in the same place, in the same years.")
body("I bring this one up last because it is the one piece of evidence in the whole book that I cannot dismiss as ancient rumor. There are people alive today who would tell you, plainly and honestly, two different histories of that stretch of years, and neither of them would be lying.")

# =====================================================================
heading1("What I Am Actually Afraid Of")
entry_date("Ninth night, and I think the last of these for now")
body("I have written seven nights of another man's teaching because it was easier than writing my own fear directly. I will do that now, plainly, and then I think I am done with this book for a while.")
body("Every account above has one thing in common that I have not said outright. Whatever else caused it, something violated the order Akatosh keeps — the plain, unglamorous work of one thing following another — badly enough that it gave way.")
body("Mehrunes Dagon is tearing open gates into this world right now, tonight, while I am writing this by a poor candle on a borrowed road. I do not know if that is the same kind of violation Vivec is describing. I hope it is not. I have no way of knowing that it is not. I am a priest who has read one difficult book three times, not a scholar, and I do not have an answer to put at the end of this the way a proper account should have one.")
body("What I have is this. Kvatch is gone. The gates are still opening. And somewhere underneath all of it, I keep thinking about a mountain in Skyrim where a debt was moved forward instead of paid, waiting for whoever comes after to settle it.")
body("My friend is asleep by the fire while I write this. Still no name given, still watching the road for me without being asked. Whatever comes next, I am glad it is not a road I am walking alone.")
attribution("-- Martin")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "What I Understand of Broken Time.docx")
doc.save(out_path)
print("SAVED:", out_path)
