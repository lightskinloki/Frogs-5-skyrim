# -*- coding: utf-8 -*-
"""Builds 'The Scourge and the Cure' -- a Tier-4 school-of-magic testament on
Destruction, in the voice of Vexis Velruan (corrected from the source's
garbled "Vexes Velron"), a temple healer expelled for turning destructive
magic against disease, who later discovered a genuine cure and destroyed
himself pursuing it. Single voice throughout, chronological through his own
descent -- NOT a professor's treatise, NOT a student's private notebook, but
a testament he intended to be read (source material's own "my dearest
reader" framing supports this). Content verified against real lore before
drafting: Vexis Velruan (temple healer, Red Fever), the Thrassian Plague
(Sload of Thras), the Knahaten Flu (Black Marsh, 2E 560), Uriel Septim VI's
death (fell from horse, 3E 313), Corprus/Dagoth Ur -- all in original prose,
not the source video's phrasing. Styled .docx, blood-red/ash palette.
Run: python _build_destruction.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Destruction: blood-red / ash-black, a testament not a lecture) ----
BLOOD   = RGBColor(0x5E, 0x14, 0x14)   # deep blood-red (headings)
EMBER   = RGBColor(0x7A, 0x2E, 0x1E)   # ember-rust (subheads)
ASH     = RGBColor(0x6E, 0x62, 0x5C)   # ash-grey accent (quote ink)
INK     = RGBColor(0x1C, 0x1C, 0x1E)
AMBER   = RGBColor(0x6B, 0x4A, 0x1E)
HEADBG  = "E8D4D0"
MARGBG  = "EAE4DC"
BORDER  = "C4A8A0"
ACCENT  = "7A2E1E"
AMBERBAR = "6B4A1E"
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = BLOOD
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = ASH
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=ASH, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def archivist_note(label, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=AMBERBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=AMBERBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = AMBER
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = AMBER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=BLOOD)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(90)
r = tp.add_run("THE SCOURGE AND THE CURE")
r.font.name = BODYFONT; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = BLOOD

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
rs = sub.add_run("Being the Testament of Vexis Velruan,\non the True Scope of the School of Destruction")
rs.font.name = BODYFONT; rs.font.size = Pt(12.5); rs.font.italic = True; rs.font.color.rgb = EMBER

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(36)
ra = attrib.add_run("Vexis Velruan")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.italic = True; ra.font.color.rgb = BLOOD

doc.add_page_break()

archivist_note("Archivist's Note -- College of Winterhold",
    "Vexis Velruan's papers were scattered and largely destroyed following his commitment to the Shivering Isles; what follows is reconstructed from the fragments that survived transcription by colleagues who could not, in the end, agree on what to make of him. We include it because his central clinical claim -- that certain diseases restoration magic cannot touch will yield to destruction magic turned against them -- has since been independently confirmed by healers who never read his work and did not know his name. We do not include it as an endorsement of his methods. -- Recordkeeper's Office"
)

# =====================================================================
# I. THE FIRST EXPERIMENT
# =====================================================================
heading1("I.  The First Experiment")
body("My dearest reader, I do not know who you will be, only that you are reading this because someone judged my words worth preserving after everything else was judged worth burning. So let me begin exactly where I would begin if you were sitting across from me, before I have had time to make it sound more dignified than it was.")
body("My eyes went first. I had not meant to aim so high -- I had meant, that first night, only to test whether a Destruction working could be turned inward without killing the caster outright, a modest question, a scholar's question -- but the spell ran hotter than I judged it, and I felt my own sight boil in its sockets before I understood what was happening to me. What came down my face afterward was not tears. I want to be precise about that, because precision is the one thing I have left that this work has not taken from me. Tears do not leave a mark on skin. What ran down my face left a mark.")
body("I broke my own hand against a door frame moments later, in a panic I am not ashamed to admit to you now, though I was ashamed of it then. It shattered the way glass shatters, all at once and completely, because the same working that had found my eyes had already reached my wrist without my noticing, and flesh that has been made brittle enough does not bend before it breaks. My legs went the same way before the night was through. I will not describe that part again. Once was enough to write it and I expect it will be enough for you to read it.")
body("Here is the thing I need you to understand, and I need you to sit with it rather than flinch past it the way every colleague I have shown this to has flinched past it: I would do it again. Not out of any love of pain, whatever they will tell you about me later. I did it because on the far side of that agony I understood something about the flesh and its frailty that no Restoration text in any library on this continent has ever managed to teach me, and I have come to believe that understanding was worth the price, even now, even having paid a good deal more of it since.")

# =====================================================================
# II. THE FIVE FACES OF DESTRUCTION
# =====================================================================
heading1("II.  The Five Faces of Destruction")
body("Before I tell you why I did what I did, I want to tell you what I actually believe about my school, because everything after this chapter depends on the reader holding it clearly, and I have found that almost no one outside serious practice holds it clearly at all.")
body("Destruction is taught to novices as the school of maximum damage in minimum time, and I will grant that this is not exactly false. It is, however, a coward's summary -- true the way a single brushstroke is true of a painting. A colleague of mine, an Illusionist by training who nonetheless felt qualified to lecture battlemages on my own discipline, once put it in precisely those terms: that Destruction is simply the art of hurting things quickly. He was not wrong about what most of his audience would go on to practice. He was wrong about what the school actually is, and I do not think he ever noticed the difference, because noticing it was never going to be useful to him.")
body("I count five faces to this discipline, not one. Elemental damage -- fire, frost, shock, and the less-discussed fourth member of that family, poison, all doing violence to matter by direct assault. Draining -- the theft of a creature's own vital energies rather than the imposition of an outside force, a quieter violence and in some ways a crueler one. Vulnerability -- not damage itself but the unmaking of a target's resistance to damage, elemental or otherwise, which includes, though almost no one includes it when they list this category, vulnerability to disease. Disintegration -- the outright unmaking of matter, armor and bodies alike, at the extreme end of what shock magic in particular is capable of. And the fifth face, the one this whole testament exists to argue for, because I have never found another mage willing to say it plainly: disease itself. Every plague that has ever swept this continent is a working of Destruction, whether or not the hand that cast it thought of itself as a mage at all.")
lore_table(
    ["Face", "What it does"],
    [["Elemental", "Fire, frost, shock, poison -- direct assault on matter by an outside force."],
     ["Draining", "Theft of a target's own vital energies rather than an external blow."],
     ["Vulnerability", "Unmaking a target's resistance -- to the elements, to injury, or to disease."],
     ["Disintegration", "The outright unmaking of matter -- armor, flesh, at the furthest reach of shock."],
     ["Disease", "The fifth face, and the one this testament is written to argue for."]],
    [1.6, 4.9])
body("A student who limits their study to the first and fourth of these -- to the fireball and the lightning bolt, the two categories every library shelf is already crowded with -- has learned perhaps a third of what this school actually contains, and has learned none of the third that I have come to believe matters most.")

# =====================================================================
# III. THE PLAGUES OF TAMRIEL
# =====================================================================
heading1("III.  The Plagues of Tamriel")
body("If the mainstream held my fifth face to be true, our histories would read very differently than they do, and I want to walk through three of them here so the reader understands I am not inventing a category to suit my own argument.")
body("Some six hundred years into the First Era, the Sload of Thras -- a race scholars still describe, not unkindly, as more slug than man -- unleashed what history now calls the Thrassian Plague upon the Iliac Bay, seeding it through infected creatures sent onto the coastline while their own people moved in behind the chaos to sack what panic had left undefended. Half the Bay's population is thought to have died of it, and by some estimates half of Tamriel besides. It took an alliance of nations who agreed on almost nothing else to end it -- the All Flags Navy, sailing to Thras itself under Bendu Olo's banner, that put the coral kingdoms under the sea by main force of arms and, I am told, considerable magic besides. No restoration ever cured that plague. Only its authors' deaths did.")
body("Centuries later and a continent away, something history calls the Knahaten Flu tore out of Black Marsh and did not stop until it had reached High Rock. One account -- and I want to be honest that it is one account among several, not a settled fact -- holds that Argonian shamans loosed it deliberately against the outsiders who had been treating their homeland as conquerable territory, and that their own people's near-total immunity to it was no accident. Whole peoples were driven to extinction in the course of it. Whatever its true origin, no priest's blessing ever slowed it. It burned itself out on its own schedule, decades after it began, having taken exactly as much as it intended to take.")
body("And in Morrowind, a disease older and stranger than either: Corprus, born from the crater where a mortal turned demigod had made himself something else entirely, warping and diseasing everything the resulting weather-thing touched. I have read that the island's own four-thousand-year-old sorcerer-king, for all his wisdom, met his limit against it -- that even the hero who eventually broke the demigod's power could strip away its madness and its worst deformities, and stopped there, the disease itself left standing. A being of that age and power found the edge of what restoration can do. I think that edge is worth someone saying plainly, and I have decided it may as well be me.")

# =====================================================================
# IV. WHAT RESTORATION CANNOT DO
# =====================================================================
heading1("IV.  What Restoration Cannot Do")
body("I trained as a healer before anyone called me anything worse, and I want that understood before I say what I am about to say, because I do not say it as an enemy of that school. I say it as someone who spent years inside it and eventually could not keep lying to myself about what I was watching.")
body("Restoration is not a guaranteed cure for anything. I have sat with dying children whose parents were told, gently and honestly, that no blessing and no potion in any temple's stores would save them, and watched those children die on schedule regardless. I have read the letters of healers working through plague season in cities I will not name, men and women who tried every remedy their training gave them -- incense, prayer, the standard restorative workings, even the folk-cures that turn out, more often than educated mages like to admit, to hold some real merit -- and who wrote, near the end of their strength, that not one soul in their care had recovered once the symptoms fully set in. Even an Emperor is not exempt from this limit. Uriel Septim the Sixth fell from his own horse and the finest healers his throne could command could not call him back. Coin and title bought him the finest hands in the Empire. Those hands knew exactly as much as restoration itself knew, and restoration itself did not know how to give him his life back.")
body("I do not say any of this to diminish Restoration's real gifts, which are considerable and which I once served with everything I had. I say it because a school that cannot cure everything has left a gap, and I have come to believe that gap is not empty. I believe something else can stand in it. I believe I found the thing that can.")

# =====================================================================
# V. THE DISCOVERY
# =====================================================================
heading1("V.  The Discovery")
body("We had a number of patients in our sanctuary infected with the Red Fever when I made my first real attempt at what became my life's actual work -- not the fireballs I had also, dutifully, learned, but an attempt to use the destructive arts to turn a disease against itself, the way one might set a controlled burn to rob a wildfire of the fuel it would otherwise reach. My early attempts failed, and failed in ways that frightened my superiors considerably more than they frightened me. I was cast out for it before I had anything resembling proof that I was right, which is, I will admit, a reasonable decision for an institution charged with the safety of the sick to have made, even if it was not a kind one.")
body("It was only in exile, without colleagues left to restrain the pace of my own research, that I found what I had been reaching for. An infection that restoration could not touch could, I confirmed past any doubt I was still capable of holding onto, be scourged directly by the deliberate application of destructive force -- burned out the way a wound is cauterized, the disease made to suffer the very violence it had been inflicting on its host. It is a brutal cure, and I will not soften that for you. But it is a cure, where none existed before, and I have never understood why that discovery bought me exile rather than a chair at any college in Tamriel.")
body("I understand it better now than I did when I wrote that last sentence some years ago and left it standing, because I have had time since to watch what the discovery did to me, and I think I understand, at last, why the great minds who might have listened were not ready for what I was offering them.")

# =====================================================================
# VI. THE HUNGER
# =====================================================================
heading1("VI.  The Hunger")
body("I will be honest with you here in a way I have not yet managed to be honest with anyone who still calls me a colleague, assuming any of them still do.")
body("Without an institution to answer to, and without patients left to test the cure on, I began testing it on myself again. The honest reason was simpler and worse than necessity: the working that burns a disease out of living flesh feels, in the moment of its casting, like nothing else I have ever experienced, and I wanted, increasingly, to feel it again. I froze my own flesh and then burned it past the point any patient would have consented to. My skin became a record of the work in a language only I could read: scars over scars, old burns crossed by newer ones, a map with no single destination left on it, only more of the same terrain repeating. It was never enough. I want to set that down exactly as plainly as it deserves. It was never enough, and I do not think, if I am honest with this page in a way I have not always managed to be honest with myself, that anything was ever going to be enough once I had felt it the first time on purpose rather than by accident.")
body("I achieved, in those years, an understanding of Magicka that I will state without false modesty exceeded anything the Grand Masters of any guild I have ever petitioned could claim for themselves. I want that on the record because it is true and because I do not expect anyone reading this after me to believe it without my saying so directly. The price of it was everything else I had left to spend, and I spent all of it, and I do not fully know, even now, whether I would call that a fair trade or simply the only trade a man in my position was ever going to be offered.")

# =====================================================================
# VII. THE LAST ENTRY
# =====================================================================
heading1("VII.  The Last Entry")
body("They have decided, finally, where I belong, and I cannot entirely argue with the decision even as I resent the manner of it. My findings on Destruction's power to heal what Restoration cannot were consigned to the fire along with the rest of what they judged unfit to preserve -- the ordinary fate, I am told, of a madman's scribblings, and I no longer have the strength to argue that the label is unearned, whatever I still believe about the work underneath it.")
body("I am to go to the Isles. I understand this is meant as a mercy, in its way, and perhaps it is one. I find I am not as afraid of that as I once would have expected to be. A realm built for the mad and the brilliant in equal measure, ruled by a Prince who is said to value both qualities without troubling himself over which is which -- I have spent a great many years now being unable to convince anyone here that those two qualities were ever separate in me to begin with. Perhaps somewhere they will finally agree with me on that point, if nothing else.")
quote("Destruction is not the art of ending things quickly. It is the art of understanding, more completely than any other school permits, exactly what a body is and exactly what it costs to unmake one -- and I have paid that cost in a currency no library will ever be willing to shelve.")
attribution("-- Vexis Velruan, formerly of the temple at Cheydinhal")

archivist_note("Archivist's Note -- closing",
    "No further writing of Vexis Velruan's has been located past this point. What became of him inside the Shivering Isles is not recorded in any source available to this College. We reproduce his testament as recovered, corrections and gaps included, and note only that the clinical claim at its center -- that certain diseases resist Restoration but yield to controlled Destruction -- is no longer considered fringe theory by the healers who have since tested it, whatever remains disputed about the man who first proved it. -- Recordkeeper's Office"
)

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Scourge and the Cure.docx")
doc.save(out_path)
print("SAVED:", out_path)
