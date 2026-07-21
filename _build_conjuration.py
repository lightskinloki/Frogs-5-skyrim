# -*- coding: utf-8 -*-
"""Builds 'On the Far Side of the Wheel' -- a Tier-4 school-of-magic treatise
on Conjuration. Single authorial voice throughout (Magister Ithendar Vael,
Master of Conjuration, College of Winterhold -- living, teaching, confident),
matching the house pattern established by 'On the Doctrine of Essences' and
'Soul Trap Tome': one professor's argued, opinionated, personally-honest
treatise, no running second voice, no narrative plot. Brief, calm archivist
framing at open and close only. Cross-checked against 'On the Structure of
Magical Law' for somatic-component consistency (Conjuration's natural
affinities: COMMAND, BIND, PROJECT, ATTUNE) and against verified lore
(Corvus & Calani Direnni; the Aurbis/Wheel cosmology). Styled .docx,
violet/aubergine palette. Run: python _build_conjuration.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (Conjuration: violet/aubergine) ----
VIOLET  = RGBColor(0x3E, 0x1E, 0x4A)
PLUM    = RGBColor(0x5A, 0x2E, 0x64)
LILAC   = RGBColor(0x7A, 0x5E, 0x82)
INK     = RGBColor(0x1C, 0x1C, 0x1E)
STEEL   = RGBColor(0x3A, 0x42, 0x48)
HEADBG  = "E6DCE9"
MARGBG  = "E4E7E8"
BORDER  = "B4A0BC"
ACCENT  = "6E4478"
STEELBAR = "3A4248"
BODYFONT = "Cambria"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = VIOLET
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = PLUM
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = LILAC
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=LILAC, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def archivist_note(label, text):
    """A brief, calm callout -- used only at open and close. Not a running voice."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=STEELBAR, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=STEELBAR, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = STEEL
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = STEEL
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=VIOLET)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(100)
r = tp.add_run("ON THE FAR SIDE OF THE WHEEL")
r.font.name = BODYFONT; r.font.size = Pt(29); r.font.bold = True; r.font.color.rgb = VIOLET

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(16)
rs = sub.add_run("Being a Treatise on the True Reach\nof the School of Conjuration")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = PLUM

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(40)
ra = attrib.add_run("Magister Ithendar Vael\nMaster of Conjuration, College of Winterhold\n4E 201")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.italic = True; ra.font.color.rgb = VIOLET

doc.add_page_break()

archivist_note("Archivist's Note -- College of Winterhold",
    "This treatise is assigned to Adept-standing students of Conjuration and above. As with all instruction in this school, students are reminded that summoning practice above the Familiar tier is subject to standard disclosure requirements under the College's agreement with the Vigilants of Stendarr -- a routine administrative matter, not a comment on this text or its author. -- Recordkeeper's Office"
)

# =====================================================================
# I. THE ONE SCHOOL WITH THE REACH
# =====================================================================
heading1("I.  The One School With the Reach")
body("Destruction, Alteration, and Restoration are, whatever their practitioners will tell you, three dialects of the same argument: each negotiates with a law already resident inside Mundus, a fragment of some elder god's own settled bargain with the shape of the world. Illusion reaches further than that only in appearance -- it touches nothing but the mortal mind that receives it, a single skull's worth of territory, however loudly its masters insist otherwise.")
body("There is one school, and one only, whose masters routinely do business with something that was never a citizen of Mundus at all, has no obligation to its laws, and in most cases has never set foot here before the moment you called it. I do not say this to boast on behalf of my discipline, although I will not pretend the boast is unearned. I say it because every apprentice who walks into their first lecture on Conjuration is handed a warning before they are handed a spell, and no apprentice in any other school receives the same courtesy on their first day, and that difference alone should tell a careful mind everything it needs to know about what separates us from the rest.")

# =====================================================================
# II. THE WHEEL, PROPERLY UNDERSTOOD
# =====================================================================
heading1("II.  The Wheel, Properly Understood")
body("No mage may practice this school with any real competence without first holding the shape of the whole cosmos in their mind, and I have found that even senior students of other disciplines cannot always draw it for me when asked. So I will draw it once, plainly, and ask the reader to hold it.")
body("Picture a wheel. At its hub sits Mundus, the mortal plane, small and central and, by any honest accounting, the least remarkable point on the whole structure. Eight spokes run outward from that hub to the wheel's rim -- the Earthbones, the portions of the elder gods who gave themselves into the world's foundation rather than abandon the labor unfinished, the settled laws that every Destruction mage and Restoration mage and Alteration mage spends their whole career negotiating with, never once looking past them to see what they are actually spokes of.")
body("Between those eight spokes lie gaps, and the gaps are not empty in the way a student's first lecture makes them sound. Each gap is a realm entire, ruled by a Prince of Oblivion who answers to no spoke and no hub, indifferent or hostile to Mundus in equal and unpredictable measure. And past the rim itself, past all sixteen realms crowded into those eight gaps, lies Aetherius -- the star-fire the Magna Ge fled into when they judged the whole project of Mundus not worth finishing, the source from which every mage of every school draws the raw Magicka they spend without ever asking where it came from.")
quote("The Spokes are the eight components of chaos, as yet solidified by the law of time.")
body("Every other school works entirely within the hub, however far its effects may seem to reach. A conjurer alone routinely reaches into the gaps. A conjurer of sufficient daring, on rare and carefully documented occasions, has touched the rim itself. I do not know a mage alive who has gone further than that and returned able to describe it. I would be lying if I said I have never wondered.")

# =====================================================================
# III. THE DIRENNI TORCH
# =====================================================================
heading1("III.  The Direnni Torch")
body("Credit for lighting this particular darkness belongs, as it happens, not to any Imperial or Breton college but to two cousins of an Altmeri clan whose name still opens doors across the Iliac Bay for reasons that have nothing to do with magic. Corvus and Calani Direnni were the first of their line, and very possibly the first of any line, to press deliberately into the gaps between the spokes rather than stumble into them by accident, and their precise binding chants are still recited by conjurers today who have never troubled to learn whose voice first shaped them.")
body("Their early results were, by any measure, a triumph. Nonbelligerent atronachs took up positions as protectors and occasional familiars to Clan Direnni's own household. Even the imp -- naturally mischievous, by every account I have read, and no easier to manage than a mortal child with a talent for lying -- was coerced into tolerably good behavior. For a generation, it must have seemed the Direnni had found a servant class more loyal than any mortal retainer, because a bound Daedra, properly bound, cannot simply decide one morning to seek better wages elsewhere.")
body("The triumph did not hold. Late in the First Era, Direnni acolytes who had grown confident on a diet of atronachs and imps began attempting the same cajolery against the Greater Daedra -- and where the most skillful among them still succeeded, reasonably, some did not, and the price of a skilled conjurer's failure against a Greater Daedra is not measured in the coin an imp might have cost. The portal those early failures widened has never fully closed since. Every mage practicing this school today inherits both halves of the Direnni legacy in the same breath: the torch, and the darkness it occasionally illuminates too well.")

# =====================================================================
# IV. THE SUMMONER'S FOUR TERMS
# =====================================================================
heading1("IV.  The Summoner's Four Terms")
body("Master Imedril of Artaeum, whose account of the somatic components I consider the single most useful text produced by any college in the last five centuries, correctly identifies Conjuration's natural fluency as resting in four terms: COMMAND, BIND, PROJECT, and ATTUNE. I want to take those four seriously in a way I do not believe Imedril himself, working from outside the discipline, was positioned to.")
body("Every other school that leans on COMMAND is forcing a law to comply. We are forcing a mind to comply -- a mind with its own history, its own grievances, very possibly its own opinion of mortals formed across centuries we cannot begin to imagine. PROJECT, for us, is not merely the distance a spell travels through open air; it is the distance between one plane of existence and another, a gap no other school's PROJECT has ever been asked to cross. BIND is the difference between a servant and a catastrophe wearing a servant's shape, and I do not use that phrase lightly. And ATTUNE -- ATTUNE is the term I have saved for last because it is the one my colleagues in other schools consistently underestimate, and it is the one I take most seriously in my own practice.")
lore_table(
    ["Component (per Imedril)", "What it costs a conjurer, specifically"],
    [["COMMAND", "Not a law's compliance, but a mind's -- a mind that may spend the entire duration of the binding calculating how to make compliance cost you everything."],
     ["BIND", "The single term separating a servant from a catastrophe temporarily wearing a servant's shape. A binding is a promise the Daedra did not choose to make."],
     ["PROJECT", "Reach across a boundary no other school's PROJECT is asked to cross -- not distance through air, but distance between one plane of existence and the next."],
     ["ATTUNE", "Partial continuity with a will that is not yours, has never been mortal, and does not stop thinking once the spell ends and you believe the conversation is over."]],
    [2.1, 4.4])
body("Imedril's own text treats these four with a scholar's remove -- descriptions, affinities, worked examples, written by a man working from outside our discipline. What I have tried to do above is tell you plainly, from inside it, what each one is actually asking of the person casting it.")

# =====================================================================
# V. THE INTERWOVEN SPELL
# =====================================================================
heading1("V.  The Interwoven Spell")
body("For most of this school's early history, summoning and binding were two separate acts, cast in sequence, and the gap between them was where conjurers died. A mage who successfully called something across from the gaps between the spokes and then miscast the binding that followed it -- or cast it a heartbeat too slowly -- was left facing an unbound, newly-arrived, and frequently furious immortal with no leash on it whatsoever. The historical record of this period is not generous to the mages who made that particular mistake.")
body("The innovation that changed this, by every account I trust, belongs to a conjurer of the following age who saw the problem plainly enough to solve it: interweave the two spells into one, so that summoning and binding manifest in the same instant, and nothing that arrives ever arrives unleashed. It is not an exaggeration to say this single insight is the reason apprentices today can attempt a Bound Dagger in their first year of study rather than their last. The Archivists' own tables bear this out precisely -- Conjure Familiar at the Novice tier is COMMAND alone, a raw calling-forth with no binding required because nothing dangerous enough to need one has yet been summoned; but the moment a mage reaches for a Flame Atronach or a Reanimated corpse, COMMAND arrives already fused to BIND, inseparable, exactly as the innovation intended.")
body("I want to be plain about what this innovation did and did not accomplish, because I have heard it mischaracterized by mages who ought to know better. It did not make Conjuration safe. It made the single most common historical cause of a conjurer's death considerably less common. Those are not the same claim, and a mage who confuses them has already made the first mistake that ends conjurers.")

# =====================================================================
# VI. ON THE LETTER AND THE SPIRIT OF A BINDING
# =====================================================================
heading1("VI.  On the Letter and the Spirit of a Binding")
body("I want to offer a case from the College's own record rather than an abstraction, since abstractions rarely land as hard as a specific failure. I have changed enough particulars that no living family need recognize it, though the lesson is unaltered.")
body("A journeyman conjurer, confident past his actual skill, once bound a minor Daedra to fetch him a rare component he could not source through ordinary trade. His binding was, by the letter, flawless -- he specified the item, the timeframe, the manner of delivery, down to the hour. He did not specify that the component should be fetched without cost to himself, because it did not occur to him that a binding needed defending against a bargain he had not offered. The Daedra, bound precisely as instructed and not one clause further, delivered the component exactly as specified -- purchased, the journeyman later learned, with a promise made in his own name to a second, unbound entity entirely, one he had never met and had no leverage over, and who considered the debt very much still open.")
body("The binding was honored to the letter in every particular. That is precisely the danger. A Daedra does not need to break your binding to ruin you. It needs only to fulfill it exactly as written, and to have understood, long before you did, everything your wording left open.")

# =====================================================================
# VII. ATTUNE, OR THE PRICE OF UNDERSTANDING ANOTHER MIND
# =====================================================================
heading1("VII.  ATTUNE, or the Price of Understanding Another Mind")
body("I said earlier that ATTUNE is the term I take most seriously in my own practice, and I owe the reader an honest account of why, rather than a scholar's shrug.")
body("To summon competently, in the fullest sense the word deserves, is not merely to force a mind across a boundary and cage it. It is, for the duration of the working, to become partially continuous with that mind -- to perceive, faintly, along the same channel the binding runs through, something of how it perceives you. I have summoned Daedra of middling rank several hundred times across my career, by careful count, and I can report that the sensation does not diminish with practice. If anything, familiarity sharpens it. A being that has existed since before Mundus took its current shape does not experience several hundred conversations with the same mortal mage as several hundred separate events. It experiences them as one long conversation, conducted at a pace that suits its own patience rather than mine, and I have begun, over the years, to suspect that what I call my many summonings, it may call something closer to a single sustained acquaintance -- one it has been quietly shaping toward some preference of its own the entire time, a preference I have had no way of detecting from my side of the exchange, because I am the one whose mind moves at mortal speed.")
body("I do not know what that preference is. I am not certain the question has an answer a mortal mind is equipped to receive even if I asked it directly. I raise it here, plainly, because a treatise that hides its author's own uncertainty from its reader has failed at the one obligation scholarship actually owes. I have not developed adequate theory to resolve this. I note it because I believe it is true, and because I would rather admit what I do not know than omit it from a text I intend to be honest.")

# =====================================================================
# VIII. WHY I CONTINUE
# =====================================================================
heading1("VIII.  Why I Continue")
body("A student nearing the end of this course has, on more than one occasion, asked me directly why I have not simply stopped -- why a mage who writes as plainly as I have about the cost of this practice keeps practicing it. I owe the reader the same honesty about my reasons that I have tried to give about the risk.")
body("Every Destruction mage in this city can tell you how fire behaves. Every Restoration mage can tell you how flesh mends. I am one of a small number of mortals currently walking under this sky who has held a genuine conversation, however constrained, however brief, with something that watched Mundus take its current shape and considered the whole affair a matter of only passing interest. That is not a comfortable thing to know. It is also, I have come to believe, worth more than my own comfort. A mortal who has never spoken with anything older than himself mistakes the hub of the wheel for the whole of the wheel, and lives his entire life believing the smallest room in the house is the house itself. I would rather carry what I have learned, at whatever cost it continues to extract from me, than go back to believing that.")
quote("A door, once understood, does not become a wall again merely because a man decides he would prefer it were one.")
attribution("-- Magister Ithendar Vael, closing this treatise")

archivist_note("Archivist's Note -- closing",
    "This text has been reviewed by the College's Conjuration faculty and is consistent with current teaching practice. Magister Vael's Chapter VII and VIII reflect his own professional judgment; students are reminded that ATTUNE-heavy practice of the kind he describes is not required coursework and should not be attempted without direct faculty supervision. -- Recordkeeper's Office"
)

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "On the Far Side of the Wheel.docx")
doc.save(out_path)
print("SAVED:", out_path)
