# -*- coding: utf-8 -*-
"""Builds 'The Way of the Voice' — a contemplative lore book on the Greybeards,
written in a Greybeard's own hand. Styled .docx, storm/slate palette."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (High Hrothgar: slate, storm, cold sky) ----
SLATE   = RGBColor(0x2E, 0x3A, 0x44)   # deep storm-slate (headings)
STORM   = RGBColor(0x47, 0x5A, 0x6B)   # weathered blue-grey (subheads)
INK     = RGBColor(0x20, 0x22, 0x24)   # near-black body ink
ASHTXT  = RGBColor(0x4A, 0x4E, 0x52)   # grey marginalia ink
HEADBG  = "DDE3E7"   # cool stone table header
MARGBG  = "ECEEF0"   # marginalia bg (cold parchment)
BORDER  = "A9B4BC"
ACCENT  = "5B6E7D"   # storm accent for rules/borders
BODYFONT = "Cambria"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = BODYFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.14

# ---- page = US Letter, 1in margins ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
CONTENT_W = 6.5

# ---- helpers ----
def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _set_cell_borders(cell, color=BORDER, sz=4, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _cell_text(cell, text, bold=False, italic=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(17); r.font.bold = True; r.font.color.rgb = SLATE
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = STORM
    return p

def body(text, italic=False, size=11, color=INK, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(size); r.font.italic = italic; r.font.color.rgb = color
    return p

def attribution(text):
    return body(text, italic=True, size=10.5, color=ASHTXT, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def epigraph(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = BODYFONT; r.font.size = Pt(11.5); r.font.italic = True; r.font.color.rgb = STORM
    return p

def marginalia(label, text):
    """A shaded callout with a thick storm-blue left border — a later hand."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(CONTENT_W)
    _set_cell_bg(cell, MARGBG)
    _set_cell_borders(cell, color=ACCENT, sz=4, sides=("top","bottom","right"))
    _set_cell_borders(cell, color=ACCENT, sz=22, sides=("left",))
    cell.text = ""
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run(label.upper())
    lr.font.name = BODYFONT; lr.font.size = Pt(8.5); lr.font.bold = True
    lr.font.color.rgb = SLATE
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(text)
    br.font.name = BODYFONT; br.font.size = Pt(10.5); br.font.italic = True; br.font.color.rgb = ASHTXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def lore_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Inches(widths[j])
        _set_cell_bg(c, HEADBG); _set_cell_borders(c)
        _cell_text(c, h, bold=True, size=9.5, color=SLATE)
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Inches(widths[j])
            _set_cell_borders(c)
            _cell_text(c, val, size=9.5)
    return t

def spacer(pts=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

# =====================================================================
# TITLE PAGE
# =====================================================================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(120)
r = tp.add_run("THE WAY OF THE VOICE")
r.font.name = BODYFONT; r.font.size = Pt(32); r.font.bold = True; r.font.color.rgb = SLATE

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(18)
rs = sub.add_run("Being a Meditation upon the Masters of High Hrothgar,\nWho Hold the Greatest Power in Skyrim and Will Not Use It")
rs.font.name = BODYFONT; rs.font.size = Pt(13); rs.font.italic = True; rs.font.color.rgb = STORM

attrib = doc.add_paragraph(); attrib.alignment = WD_ALIGN_PARAGRAPH.CENTER
attrib.paragraph_format.space_before = Pt(60)
ra = attrib.add_run("Set down in silence by one of the Greybeards,\nfor the hand that keeps the vigil after mine.")
ra.font.name = BODYFONT; ra.font.size = Pt(12); ra.font.color.rgb = INK

warn = doc.add_paragraph(); warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
warn.paragraph_format.space_before = Pt(48)
rw = warn.add_run("We do not speak. So I have written.\nRead it as you would hear it: slowly, and only once aloud in your heart.")
rw.font.name = BODYFONT; rw.font.size = Pt(11); rw.font.italic = True; rw.font.color.rgb = ASHTXT

doc.add_page_break()

# =====================================================================
# I. WHY I WRITE AT ALL
# =====================================================================
heading1("I.  Why I Write at All")
body("A Greybeard does not speak, save in great need, and a Greybeard who speaks at his ease is no Greybeard at all. You will know this before you can read these words, for it will have been the first thing taught you on the mountain, and the hardest. So you will wonder why a master of the order has taken up the pen, when the pen is itself a small surrender — a way to put out into the world what the mouth is forbidden to send.")
body("I write because silence is not ignorance, and our restraint is not the same thing as forgetting. We keep our voices still so that the world may stand; we keep our memory whole so that the stillness means something. There must be one record, in a plain hand, of who we are and why we sit on this peak with the power to unmake kingdoms folded quietly in our throats. Not a record of how the Voice is shaped — that is taught mouth to ear, master to student, and belongs to no page. A record of why we choose, every single day, not to use it.")
body("If you hold this book, the vigil is yours now, or soon will be. Read it slowly. Then keep your silence better than I kept mine, for I have just broken mine to make it.")
attribution("— a Master of High Hrothgar")

# =====================================================================
# II. THE BEAST-GODS AND THE FIRST REVERENCE
# =====================================================================
heading1("II.  The Beast-Gods and the First Reverence")
body("Before there was a Skyrim there was Atmora, and before there was an Empire there were the old Nords who crossed the ice from it. They came with their gods, and their gods wore no human faces. The divine, to the first Nords, looked out through the eyes of beasts. There was the Hawk and the Wolf, the Snake and the Moth, the Owl, the Whale, the Bear, and the Fox — and over all of them, greatest and most terrible, the Dragon.")
body("It is easy now to call this a child's faith and pass it by. Do not. The first Nords were not wrong to see the divine in the beast; they were only too willing to kneel to it. They revered the Dragons, and built for them great temples of stone, and bowed in grand displays of obedience — and the Dragons, who are pride given wings, accepted the worship as their due. This is the root of everything that came after: a people who looked upon raw power and called it holy, and a power that looked upon worship and called it permission.")
epigraph("The beast-faces of the old gods, as the first Nords reckoned them —")
lore_table(
    ["The face", "As the old Nords revered it"],
    [["The Dragon", "Greatest of all; God-king above gods and men alike. Worship was owed, and taken."],
     ["The Hawk", "The wind and the heights; later remembered in Kyne, the Mother."],
     ["The Wolf", "The hunt and the loyal pack."],
     ["The Snake", "The cunning and the hidden turn."],
     ["The Moth", "The quiet wisdom; the reader of what is written in light."],
     ["The Owl, Whale, Bear, Fox", "The lesser faces — watcher, deep, strength, and trickster — each honored in its season."]],
    [1.6, 4.9])
spacer()

# =====================================================================
# III. THE DRAGON WAR AND THE GIFT
# =====================================================================
heading1("III.  The Dragon War, and the Gift That Ended It")
body("The Dragons did not rule alone. They raised up men to rule for them — the Dragon Priests, who began as servants of gods and ended as tyrants over their own kind. Power that flows downhill from a god corrupts the vessel it passes through, and the priests were corrupted utterly. They ground their fellow men under an iron rule, and when at last the people rose against it, the Dragons came down from the sky and answered the rebellion with slaughter. Thousands of Nords were butchered. The war was swift, and it was very nearly the end of us.")
body("It was not the end, because not every Dragon was content to be worshipped. Some pitied the small, soft creatures dying in their thousands. Chief among these was Paarthurnax, who was the younger brother of Alduin, the World-Eater himself who led the Dragons. Paarthurnax turned from his own kind and gave to men the one thing that could stand against a Dragon: the Voice. He taught mortals the Thu'um — to shape the dragon-tongue, to speak as the Dragons speak, so that the word of a man might strike with the weight of the word of a god.")
body("With that gift the tide turned. The Dragons were driven off and slain in their multitudes, and Alduin himself was cast forward out of time, banished by the cunning of an Elder Scroll to be dealt with in an age not yet come. When it was over, Paarthurnax did not take a crown for his mercy. He withdrew to the summit of the Throat of the World, the highest peak in all Tamriel, and there he has remained — teacher to those few mortals worthy of being taught. That summit is our mountain. Everything we are descends from that one Dragon's choice to pity us.")
marginalia("a later hand, in the margin",
    "The student should hold both halves of this and not let go of either. The Voice was the salvation of men — and the Voice was first the weapon of the gods who nearly destroyed us. The same power saved us and damned us. That is not a contradiction to be resolved. It is the reason our order exists.")

# =====================================================================
# IV. JURGEN WINDCALLER
# =====================================================================
heading1("IV.  Jurgen Windcaller, and the Silence")
body("For an age men used the gift freely, and called it glory. They were called tongues — mortals who carried the Voice — and the mightiest of them could level walls and scatter armies with a shout. The first Empire of the Nords was built on such voices, hungry for land, and in the First Era it reached out to take the country across the Velothi mountains, the land we now call Morrowind. There, on the foothills of Red Mountain, the joined strength of the Chimer and the Dwemer met them and broke them utterly. The Nords were driven from the land, and the dead were past counting.")
body("Among the survivors was the greatest tongue who has ever lived: Jurgen Windcaller. He had been a warrior and a warlord, and he had shouted men into their graves without a second thought. The defeat undid him. He could not reconcile it. How could voices so strong have failed so completely? How could so many be dead, with all that power in their throats? He withdrew from the world and meditated upon the question for seven years, and at the end of them he had his answer, and it changed everything.")
body("He concluded that the gods had punished the Nords — that the Voice was a gift of Kyne and the gods, given to honor them, and that men had profaned it by turning it to conquest and the seizing of foreign lands. The Dragons had once been punished for the tyrannous use of their power; now men had been punished for the same arrogance. The Voice, Jurgen taught, is not the property of those who can shape it. It is a holy thing, owed back to the gods in reverence, and to spend it on the victories of men is blasphemy.")
heading2("The Seventeen, and the Three Days")
body("Many were moved by his teaching. More were enraged by it, for he asked the strong to lay down the very thing that made them strong. Seventeen masters of the Voice came to challenge him, and they assailed him with the full fury of their shouts. Jurgen would not answer them in kind. He stood, and he did not rise to a single provocation, and for three days he simply absorbed everything they could send against him — swallowed their shouts whole until the seventeen had spent themselves and stood before him with parched throats and broken voices, beaten not by violence but by his refusal of it. They acknowledged his wisdom. The Way of the Voice was founded that day.")
body("Consider what he proved. The ultimate display of mastery over the Voice was to choose, in the face of every reason to shout, not to shout. For a man whose whole power lives in his speech, silence is the highest discipline there is. That is the paradox we keep. We are not weak. We are the opposite of weak, holding still.")
attribution("— so it was taught to me, as I now teach it to you")

# =====================================================================
# V. THE WAY OF THE VOICE
# =====================================================================
heading1("V.  The Way of the Voice")
body("From Jurgen's wisdom the order took its shape, and its rule is simple to say and a lifetime to keep. The Voice is to be used in reverence to Kyne and the gods, and not in the affairs of men and mer. We do not take sides in wars. We do not raise up kings or pull them down. We do not lend our power to any cause, however just it seems, because the moment we judge a cause worthy of the Voice we have set ourselves above the gods who gave it, and that is the very arrogance that nearly ended our people twice over.")
body("So we keep silence — not as a vow of suffering, but as the daily practice of the discipline Jurgen proved. We speak to one another and to the world in writing, in gesture, in the patient sign that costs no breath. A student spends years learning to still the Voice before he is ever taught to raise it, for an unstilled Voice in an untrained throat is a danger to everyone near it. We are monks before we are tongues. The shout is the least of what we teach; the silence is the whole of it.")
marginalia("a later hand, in the margin",
    "Do not mistake our peace for ease. Every master on this mountain carries within him the means to bring down an avalanche, to shatter a gate, to end a life with a word — and chooses, hour upon hour, year upon year, to swallow it. The world sleeps soundly because we do not. Remember that when the silence grows heavy on you, as it will.")

# =====================================================================
# VI. HIGH HROTHGAR AND THE SEVEN THOUSAND STEPS
# =====================================================================
heading1("VI.  High Hrothgar and the Seven Thousand Steps")
body("Our monastery is High Hrothgar, set upon the slopes of the Throat of the World, far above the goings-on of common people. We chose the height on purpose. Distance is part of the discipline; a man cannot be drawn into the quarrels of the valley if the valley is a week's hard climb below him. The breath of Kyne upon this mountain is held by the Nords to be the very circumstance of their creation, the place where life was first breathed into the folk of Skyrim, and so the climb to us is also a climb toward the beginning of all Nordkind.")
body("Any Nord may make the pilgrimage. Seven thousand steps wind up the mountain to our gates, and the climbing of them is itself the teaching — for by the time a pilgrim reaches us, the world's noise has fallen away behind and below, and what remains is breath, and cold, and the slow work of putting one foot above the last. Many who set out do not arrive. The mountain is honest in that way. It does not lie about what it costs to come near the gods.")

# =====================================================================
# VII. THE POWER WE WILL NOT USE
# =====================================================================
heading1("VII.  The Power We Will Not Use")
body("It must be understood plainly, lest the silence be mistaken for the whole truth: we are not silent because we are gentle. We are silent because we are not. The power held on this mountain could tear down empires and reorder the face of Skyrim in an afternoon, and it is precisely because this is true that we never reach for it. A power that can be used lightly will be; ours cannot be used lightly at all, and so we have chosen not to use it.")
body("When a master of this order so much as speaks aloud upon the mountain, storms gather. The sky darkens, and the people of the lower villages take warning and flee, for an avalanche may follow the merest greeting. There have been masters in our history whose Voices grew so strong that to engage the throat at all brought destruction down instantly, and they were obliged to remain wholly silent for the rest of their lives, speaking only by sign and by the written hand. This is not legend. It is the plain hazard of the gift carried to its height.")
body("It is told that when the Greybeards once spoke a greeting to a high king who came believing himself a figure of prophecy, the greeting alone reduced him to ash — not from malice, but because the Voice does not soften itself for the unready, and he was unready. And when the order announced the name of the young Tiber Septim, who would go on to forge the Empire, the speaking of his name shook the world. We do not say such names often. The world cannot afford our enthusiasm.")
marginalia("a later hand, in the margin",
    "There is a story the alchemists tell of a vial that purifies any liquid poured into it and replenishes itself forever — a miracle cure, made with snow that would not melt even in the fires of the Deadlands. They say the snow refuses the sun because a Greybeard once taught it to. I do not confirm the tale. I set it here only so the student understands the shape of what we are: a word from this mountain can teach water to forget it is water, and frost to forget the sun. Weigh that, and then weigh our silence, and you will know why we keep it.")

# =====================================================================
# VIII. THE HERALDING OF THE DRAGONBORN
# =====================================================================
heading1("VIII.  The Heralding of the Dragonborn")
body("There is one cause for which we will break our silence, and only one. When a Dragonborn comes into the world, we are the first to know it, and we proclaim it — for the Dragonborn is not a tongue who studied as we studied, but a mortal blessed by Akatosh himself, father of Dragons and chief of the divine, with the very blood and soul of a dragon. Such a one does not labor years to learn the dragon-tongue. They take it directly from the soul of a slain dragon, absorbing in an instant what costs us a lifetime. They are the gods' own answer to the arrogance of the powerful, sent into the world to crush blasphemy and set the balance right.")
body("So we listen. It is the deepest of our offices, older than the silence itself: to keep watch for the one akatosh marks, and when the mantle settles upon a mortal soul, to send word from the mountain so that all Skyrim may know the prophecy is moving. We heralded Tiber Septim. We have heralded others before him and since. To call the name of a Dragonborn is the one shout we are permitted without shame, for it is the gods' work and not the work of men.")
marginalia("a later hand, pressed harder into the page",
    "And here I must set down a thing that troubles the order, and that the next vigil-keeper must carry: the listening has gone quiet. There are signs in the world — dragons stirring that should sleep, an unease in the deep places — that a Dragonborn walks among the living even now. Yet the mantle does not ring on the mountain as the old records swear it always has. We strain to hear, and hear nothing. Either the gods have changed how they mark their chosen, or this one is hidden even from us. Do not assume the silence means there is no one. Assume only that we cannot yet hear them — and keep listening the harder for it.")

# =====================================================================
# IX. THE BLADES
# =====================================================================
heading1("IX.  The Blades, Who Do Not Understand Us")
body("You will hear, in time, of the Blades — once the Akaviri Dragonguard, sworn for an age to the service of the Dragonborn. We share with them a reverence for the Dragonborn and almost nothing else, and the gulf between us is old and bitter. They hate us for our stillness, naming it cowardice, faulting us for every war we did not win for them and every enemy we did not strike down on their behalf. We, for our part, hold them to be meddlers and would-be makers of slaughter, who have always sought to turn the Dragonborn away from wisdom and toward the sword.")
body("Understand the quarrel rightly, for you will be tempted to take a side. The Blades believe power exists to be used, and that to hold it back in the hour of need is betrayal. We believe power exists to be revered, and that to spend it in the affairs of men is the betrayal. Neither of us will be argued out of it. When a Dragonborn comes, both will reach for that soul — the Blades to make of it a weapon, we to make of it a pilgrim. That tug-of-war is as old as the Dragonborn prophecy, and it has never once been settled gently.")

# =====================================================================
# X. CLOSING
# =====================================================================
heading1("X.  In Closing")
body("So you have it now, in a plain hand: where we came from, what we learned, and why we sit upon this peak with the strength to remake the world and the discipline never to. We watched the Nords nearly destroyed by the worship of raw power, and then nearly destroyed again by the wielding of it, and we drew from those two near-endings a single conclusion — that the only safe place for the greatest power in Skyrim is in the keeping of those who have sworn not to use it.")
body("It is a hard creed. There will be hours, when the valley below us burns and the people we will not save cry out, that the silence will feel like sin. Hold it anyway. We are bound by the Way of the Voice to be idle bystanders to the wars of men, even unto the loss of our own land, because the alternative — a mountain of monks who decide which wars are worthy of a god's own power — is a horror greater than any war. Keep the vigil. Keep the silence. And listen, always, for the one the gods will send to do what we never can.")
attribution("— a Master of High Hrothgar, in the keeping of the vigil")

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lore Books")
out_path = os.path.join(out_dir, "The Way of the Voice.docx")
doc.save(out_path)
print("SAVED:", out_path)
